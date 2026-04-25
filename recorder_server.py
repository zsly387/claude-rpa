#!/usr/bin/env python3
"""
RPA Recorder Server
===================
Long-running headed Playwright process started by `rpa_manager record-start`.

IPC protocol (file-based):
  rpa_manager writes  → SESSION_DIR/cmd.json   {"action":..., "seq": N, ...}
  recorder_server writes → SESSION_DIR/result_N.json  {"success":..., "snapshot":[], ...}

Each executed action:
  - Runs in the headed (visible) Chromium window (browser steps), or runs locally for
    `api_call`, `merge_files`, `excel_write`, `word_write` (still snapshots the page after)
  - Takes a screenshot
  - Returns DOM snapshot so the LLM can pick real CSS selectors
  - Appends generated Python code to code_blocks list

On shutdown:
  - Compiles code_blocks into a full standalone Playwright script
  - Saves to SESSION_DIR/script_log.py
  - Writes SESSION_DIR/done marker
"""

import asyncio
import inspect
import json
import os
import re
import sys
import urllib.parse
from datetime import datetime
from pathlib import Path

import httpx

# ── 时间戳日志：所有 print() 自动加前缀 [HH:MM:SS] ──────────────────────────
# Timestamped logging: prepend [HH:MM:SS] to every print() call
import builtins as _builtins
_real_print = _builtins.print
def _timed_print(*args, **kwargs):  # type: ignore[override]
    _real_print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}]", *args, **kwargs)
_builtins.print = _timed_print
# ─────────────────────────────────────────────────────────────────────────────

SKILL_DIR   = Path(__file__).parent
SESSION_DIR = SKILL_DIR / "recorder_session"

POLL_INTERVAL = 0.15  # seconds

# 同一次录制会话内，同一输出文件名多次 extract_text：首次 write_text，之后追加
_EXTRACT_OUTPUT_FILES: set[str] = set()

# 本次 session 中每个 kv 文件写入了哪些字段（供 python_snippet 结构性卡口使用）
# 格式：{输出文件名: [字段名列表]}  例：{"hotel1.txt": ["民宿名字", "评分", "价格"]}
_EXTRACT_FIELD_REGISTRY: dict[str, list[str]] = {}

# 本次 session 是否已有 python_snippet 成功使用了 page.evaluate 提取列表行型数据。
# 为 True 时，后续 python_snippet 的结构性卡口（_parse_field 要求）自动豁免——
# 因为数据通道是 page.evaluate → JSON 文件，而非 extract_text → _parse_field。
_SESSION_HAS_PAGE_EVALUATE: bool = False

# 本次 session 的临时文件目录：/tmp/{task_slug}/
# 每次 record-start 由 server_main() 设置，隔离不同任务的提取文件互不干扰
_TASK_TMP_DIR: Path = Path("/tmp") / "rpa_default"


def _slugify_for_path(text: str) -> str:
    """将任务名转为安全目录名（保留中文、字母、数字、连字符，其余替换为下划线）。"""
    import re as _sre
    slug = _sre.sub(r'[^\w\u4e00-\u9fff-]', '_', text.strip())
    slug = _sre.sub(r'_+', '_', slug).strip('_')
    return slug[:48] or "task"


def _reset_extract_output_tracking() -> None:
    global _EXTRACT_OUTPUT_FILES, _EXTRACT_FIELD_REGISTRY, _VISION_SESSION, _VISION_STEPS, \
           _SESSION_HAS_PAGE_EVALUATE
    _EXTRACT_OUTPUT_FILES = set()
    _EXTRACT_FIELD_REGISTRY = {}
    _VISION_SESSION = {}
    _VISION_STEPS = []
    _SESSION_HAS_PAGE_EVALUATE = False


_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

# ── 视觉识别模型配置表 ─────────────────────────────────────────────────────────
# 两个默认模型均使用 OpenAI-compatible 格式，共用同一套调用代码。
# 视觉步骤 HTTP 超时（秒）：VL + 全页截图可能较慢，须大于 rpa_manager 对 extract_by_vision 的轮询等待
_VISION_HTTP_TIMEOUT_SEC = 300.0

_VISION_MODELS: dict[str, dict] = {
    "qwen": {
        "label":    "Qwen3-VL-Plus（阿里云百炼，多模态）",
        "model":    "qwen3-vl-plus",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "key_env":  "DASHSCOPE_API_KEY",
        "key_url":  "https://bailian.console.aliyun.com → API Key 管理",
    },
    "gemini": {
        "label":    "Gemini 3 Pro（Google AI Studio）",
        "model":    "gemini-3-pro-preview",
        "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
        "key_env":  "GOOGLE_AI_KEY",
        "key_url":  "https://aistudio.google.com/app/apikey",
    },
}

# 本次 session 的视觉识别配置（由 record-step 的 model_key + api_key 写入）
_VISION_SESSION: dict = {}   # {"model_key": "qwen", "api_key": "sk-xxx"}

# 本次 session 录制的视觉步骤（用于 record-end 时生成 vision_setup.md）
_VISION_STEPS: list[dict] = []

# 与 SKILL 中「重型 SPA 域名表」对齐（hostname 小写；等于根域或为其子域）
_HEAVY_SPA_HOST_ROOTS: tuple[str, ...] = (
    "airbnb.cn",
    "airbnb.com",
    "booking.com",
    "hotels.com",
    "agoda.com",
    "trivago.com",
    "trip.com",
    "ctrip.com",
    "fliggy.com",
    "xiaohongshu.com",
    "xhslink.com",
    "douyin.com",
    "iesdouyin.com",
    "tiktok.com",
    "instagram.com",
    "facebook.com",
    "linkedin.com",
    "twitter.com",
    "x.com",
    "maps.google.com",
    "openrice.com",
    "yelp.com",
    "shein.com",
)


def _hostname_on_heavy_spa_list(host: str) -> bool:
    h = (host or "").lower().strip()
    if not h:
        return False
    if h.startswith("shopee.") or ".shopee." in h:
        return True
    if h.startswith("expedia.") or ".expedia." in h:
        return True
    for root in _HEAVY_SPA_HOST_ROOTS:
        if h == root or h.endswith("." + root):
            return True
    return False


def _vision_keys_path() -> Path:
    p = Path.home() / ".openclaw" / "rpa" / "vision_keys.json"
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


def _load_cached_vision_key(model_key: str) -> str:
    """从本地缓存读取 API key；不存在返回空字符串。"""
    p = _vision_keys_path()
    if not p.exists():
        return ""
    try:
        return json.loads(p.read_text(encoding="utf-8")).get(model_key, "")
    except Exception:
        return ""


def _save_vision_key(model_key: str, api_key: str) -> None:
    """将 API key 持久化到本地缓存，下次录制自动复用。"""
    p = _vision_keys_path()
    try:
        data = json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}
    except Exception:
        data = {}
    data[model_key] = api_key
    p.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


async def _call_vision_api(
    image_bytes: bytes,
    fields: list[str],
    model_key: str,
    api_key: str,
) -> dict[str, str]:
    """调用视觉 LLM，从截图中提取指定字段，返回 {字段名: 值} 字典。

    Qwen 与 Gemini 均使用 OpenAI-compatible chat completions 格式：
      Qwen：  dashscope.aliyuncs.com/compatible-mode/v1/chat/completions
      Gemini：generativelanguage.googleapis.com/v1beta/openai/chat/completions
    """
    import base64 as _b64
    cfg = _VISION_MODELS[model_key]
    b64 = _b64.b64encode(image_bytes).decode()
    fields_tmpl = json.dumps({f: "" for f in fields}, ensure_ascii=False)
    prompt = (
        f"从截图中提取以下字段，只返回 JSON，不要任何解释或 markdown 代码块：\n"
        f"{fields_tmpl}\n\n"
        "规则：①只提取截图中实际可见的文字；"
        "②看不到的字段设为空字符串；"
        "③价格保留原始格式（如 ¥368/晚）。"
    )
    payload = {
        "model": cfg["model"],
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                {"type": "text", "text": prompt},
            ],
        }],
        "max_tokens": 500,
    }
    base = cfg["base_url"].rstrip("/")
    _t = httpx.Timeout(_VISION_HTTP_TIMEOUT_SEC, connect=30.0)
    async with httpx.AsyncClient(timeout=_t, verify=False) as hc:
        r = await hc.post(
            f"{base}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=payload,
        )
        r.raise_for_status()
    raw = r.json()["choices"][0]["message"]["content"].strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.MULTILINE).strip()
    return json.loads(raw)


async def _validate_vision_key(model_key: str, api_key: str) -> tuple[bool, str]:
    """用最小请求验证 API key 是否可用，返回 (ok, error_msg)。"""
    # 最小合法 PNG（1×1 白色像素）
    _TINY_PNG = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff"
        b"?\x00\x05\xfe\x02\xfe\xdc\xccY\xe7\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    try:
        await _call_vision_api(_TINY_PNG, ["test"], model_key, api_key)
        return True, ""
    except httpx.HTTPStatusError as e:
        if e.response.status_code == 401:
            return False, "API Key 无效（401 Unauthorized）"
        if e.response.status_code == 429:
            return True, ""   # 限速 = key 有效
        return False, f"HTTP {e.response.status_code}: {e.response.text[:100]}"
    except json.JSONDecodeError:
        return True, ""   # 能调通就算有效
    except Exception as e:
        return False, str(e)[:120]


# 与 _build_final_script 生成脚本中的 _EXTRACT_JS 一致。
# 裸标签选择器（如 h3、无 # . [ 空格）在存在 <main> / [role=main] 时只在该区域内匹配，
# 避免 Yahoo 等站顶栏 mega-menu 的 h3 先于正文被 slice(0,n) 取走。
_EXTRACT_JS_MIN = (
    '([s,n])=>{const r=document.querySelector("main")||document.querySelector(\'[role="main"]\');'
    'const bare=/^[a-zA-Z][a-zA-Z0-9-]*$/.test(s)&&s.indexOf("#")<0&&s.indexOf(".")<0&&'
    's.indexOf("[")<0&&s.indexOf(" ")<0;'
    'const sc=bare&&r?r:document;return Array.from(sc.querySelectorAll(s)).slice(0,n)'
    '.map(e=>(e.textContent||"").replace(/\\s+/g," ").trim()).filter(Boolean)}'
)


async def _wait_spa_ready_for_vision(
    page,
    crop_selector: str = "",
    *,
    timeout_ms: int = 45_000,
) -> None:
    """视觉截图前等待 SPA 主内容就绪，减少骨架屏、未 hydration 就截图的情况。

    顺序：domcontentloaded → 尝试 networkidle → 固定短等 → 轮询「大图已解码或正文足够」；
    轮询中偶尔 wheel 触发懒加载。若提供 crop_selector，最后再等该容器 visible。
    """
    import time as _time

    try:
        await page.wait_for_load_state("domcontentloaded", timeout=12_000)
    except Exception:
        pass
    try:
        await page.wait_for_load_state("networkidle", timeout=28_000)
    except Exception:
        pass
    await page.wait_for_timeout(1_200)

    deadline = _time.monotonic() + max(5_000, timeout_ms) / 1000.0
    poll_js = """() => {
        const imgs = Array.from(document.querySelectorAll('img'));
        for (const i of imgs) {
            if (i.complete && i.naturalWidth > 64 && i.naturalHeight > 64) return true;
        }
        const t = (document.body && document.body.innerText) || '';
        const compact = t.replace(/\\s+/g, '');
        if (compact.length > 380 && /\\d/.test(t)) {
            if (/[\\u4e00-\\u9fff]/.test(t) || compact.length > 620) return true;
        }
        return false;
    }"""
    n = 0
    while _time.monotonic() < deadline:
        try:
            if await page.evaluate(poll_js):
                await page.wait_for_timeout(700)
                break
        except Exception:
            pass
        n += 1
        if n % 5 == 0:
            try:
                await page.mouse.wheel(0, 320)
            except Exception:
                pass
        await page.wait_for_timeout(420)
    else:
        await page.wait_for_timeout(1_000)

    if crop_selector and str(crop_selector).strip():
        try:
            await page.locator(crop_selector.strip()).first.wait_for(
                state="visible", timeout=min(25_000, timeout_ms)
            )
            await page.wait_for_timeout(400)
        except Exception:
            pass


# ── Code generation helpers ──────────────────────────────────────────────────

def _step_code(step_n: int, context: str, body: list[str]) -> str:
    """Wrap body lines in a try/except block inside async def run()."""
    ind = "            "  # 12 spaces — sits inside: async with → browser → try
    lines = [
        f"{ind}# ── 步骤 {step_n}：{context}",
        f"{ind}try:",
    ]
    for b in body:
        lines.append(f"{ind}    {b}")
    lines += [
        f"{ind}except Exception:",
        f'{ind}    await page.screenshot(path="step_{step_n}_error.png")',
        f"{ind}    raise",
    ]
    return "\n".join(lines)


# params / headers 字符串值填 __ENV:YOUR_ENV_VAR_NAME__，录制与回放均从环境变量读取
_ENV_PLACEHOLDER_RE = re.compile(r"^__ENV:([A-Za-z_][A-Za-z0-9_]*)__$")


def _resolve_placeholders_for_record(v, data: dict):
    """将占位符替换为环境变量或 data['env'] 中的值（仅录制时用于真实 HTTP 请求）。"""
    if not isinstance(v, str):
        return v
    env_fallback = data.get("env") or {}
    m = _ENV_PLACEHOLDER_RE.match(v)
    if m:
        name = m.group(1)
        return os.environ.get(name, env_fallback.get(name, ""))
    return v


def _params_for_record(data: dict) -> dict:
    params = dict(data.get("params") or {})
    for k, v in list(params.items()):
        params[k] = _resolve_placeholders_for_record(v, data)
    return params


def _headers_for_record(data: dict):
    h = data.get("headers")
    if not h:
        return None
    return {k: _resolve_placeholders_for_record(v, data) for k, v in h.items()}


def _build_api_url_for_record(data: dict) -> str:
    if data.get("base_url") is not None and data.get("params") is not None:
        base = str(data["base_url"]).rstrip("?")
        q = urllib.parse.urlencode(_params_for_record(data))
        return f"{base}?{q}"
    url = data.get("url")
    if not url:
        raise ValueError("api_call 需要 url，或 base_url + params / api_call requires 'url' or 'base_url' + 'params'")
    return str(url)


from typing import Optional
def _codegen_env_value(v, env_values: Optional[dict] = None) -> str:
    """params/headers 里单个值对应的 Python 表达式。

    若 v 是 __ENV:VAR__ 占位符：
    - 若 env_values 中有该变量的真实值（用户在录制时已提供），直接写入脚本，免去 export。
    - 否则生成 os.environ.get(...)，运行时从环境变量读取。
    """
    if not isinstance(v, str):
        return repr(v)
    m = _ENV_PLACEHOLDER_RE.match(v)
    if m:
        name = m.group(1)
        if env_values and env_values.get(name):
            # User already supplied the real value — embed it directly so the
            # generated script runs without any `export` setup.
            return repr(env_values[name])
        # Use double quotes consistently so _build_final_script regex can find them.
        return f'os.environ.get("{name}", "")'
    return repr(v)


def _api_codegen_body(context: str, data: dict) -> list[str]:
    """生成 run() 内 httpx 调用代码行（不含 try/except 外壳）。"""
    method = (data.get("method") or "GET").upper()
    save_to = data.get("save_response_to")
    headers = data.get("headers")
    env_values: dict = data.get("env") or {}
    # 自动检测含下划线的主机名 → 生成 verify=False
    _url_for_check = data.get("url") or data.get("base_url") or ""
    _host_for_check = urllib.parse.urlparse(_url_for_check).hostname or ""
    _auto_no_verify = "_" in _host_for_check
    verify_ssl = data.get("verify_ssl", not _auto_no_verify)
    verify_kw = "" if verify_ssl else ", verify=False"

    lines: list[str] = []
    if data.get("base_url") is not None and data.get("params") is not None:
        parts: list[str] = []
        for k, v in data["params"].items():
            parts.append(f"{repr(k)}: {_codegen_env_value(v, env_values)}")
        lines.append("_params = {" + ", ".join(parts) + "}")
        lines.append(
            f'_api_url = {repr(str(data["base_url"]).rstrip("?"))} + "?" + urllib.parse.urlencode(_params)'
        )
    else:
        url = data.get("url")
        lines.append(f"_api_url = {repr(url)}")

    hdr_kw = ""
    if headers:
        hparts = [f"{repr(k)}: {_codegen_env_value(v, env_values)}" for k, v in headers.items()]
        lines.insert(0, "_api_headers = {" + ", ".join(hparts) + "}")
        hdr_kw = ", headers=_api_headers"

    client_kw = f'timeout=CONFIG["api_timeout"]{verify_kw}'
    if method == "GET":
        lines.append(f"async with httpx.AsyncClient({client_kw}) as _hc:")
        lines.append(f"    _r = await _hc.get(_api_url{hdr_kw})")
        lines.append("    _r.raise_for_status()")
    elif method == "POST":
        body = data.get("body")
        lines.append(f"async with httpx.AsyncClient({client_kw}) as _hc:")
        if isinstance(body, dict):
            if hdr_kw:
                lines.append(f"    _r = await _hc.post(_api_url, json={repr(body)}, headers=_api_headers)")
            else:
                lines.append(f"    _r = await _hc.post(_api_url, json={repr(body)})")
        else:
            b = body if body is not None else ""
            if hdr_kw:
                lines.append(f"    _r = await _hc.post(_api_url, content={repr(b)}, headers=_api_headers)")
            else:
                lines.append(f"    _r = await _hc.post(_api_url, content={repr(b)})")
        lines.append("    _r.raise_for_status()")
    else:
        lines.append(f"async with httpx.AsyncClient({client_kw}) as _hc:")
        lines.append(
            f"    _r = await _hc.request({repr(method)}, _api_url{hdr_kw})"
        )
        lines.append("    _r.raise_for_status()")

    if save_to:
        lines.append(f'(CONFIG["tmp_dir"] / {repr(save_to)}).write_text(_r.text, encoding="utf-8")')
        lines.append(f'print("API 响应已写入", CONFIG["tmp_dir"] / {repr(save_to)})')
    else:
        lines.append('print("API 响应长度:", len(_r.text))')

    return lines


def _resolve_file(file_str: str, base_dir: Path, fallback_dir: Optional[Path] = None) -> Path:
    """若 file_str 是绝对路径直接返回；否则先在 base_dir 查找，找不到再试 fallback_dir。"""
    p = Path(file_str)
    if p.is_absolute():
        return p
    candidate = base_dir / file_str
    if not candidate.exists() and fallback_dir is not None:
        fb = fallback_dir / file_str
        if fb.exists():
            return fb
    return candidate


def _excel_rows_from_json(spec: dict, base_dir: Path, fallback_dir: Optional[Path] = None) -> list[list]:
    """展平 JSON 文件中的嵌套数组，返回二维行列表。

    spec 格式（任选其一）：
      平铺：{"file":"x.json","outer_key":"items","fields":["f1","f2"]}
      嵌套：{"file":"x.json","outer_key":"batches","inner_key":"lines",
             "fields":["f1","f2"],"parent_fields":["batch_id"]}
    file 可为绝对路径，也可为相对于 base_dir 的文件名（找不到时 fallback 到 fallback_dir）。
    """
    import json as _json
    fpath = _resolve_file(spec["file"], base_dir, fallback_dir)
    if not fpath.exists():
        return []
    data = _json.loads(fpath.read_text(encoding="utf-8"))
    outer_key = spec.get("outer_key", "")
    inner_key = spec.get("inner_key", "")
    fields = spec.get("fields") or []
    parent_fields = spec.get("parent_fields") or []
    outer_list = data.get(outer_key, []) if outer_key else (data if isinstance(data, list) else [])
    rows: list[list] = []
    if inner_key:
        for outer_item in outer_list:
            pvals = [outer_item.get(pf) for pf in parent_fields]
            for inner_item in outer_item.get(inner_key, []):
                rows.append([inner_item.get(f) for f in fields] + pvals)
    else:
        for item in outer_list:
            rows.append([item.get(f) for f in fields] + [item.get(pf) for pf in parent_fields])
    return rows


def _excel_rows_from_excel(spec: dict, base_dir: Path, fallback_dir: Optional[Path] = None) -> list[list]:
    """从另一个 xlsx 文件的指定 sheet 读取数据行（不含表头首行）。

    spec 格式：{"file":"发票导入_本周.xlsx","sheet":"发票侧","skip_header":true}
    file 可为绝对路径，也可为相对于 base_dir 的文件名（找不到时 fallback 到 fallback_dir）。
    skip_header 默认 true，跳过第一行表头。
    """
    from openpyxl import load_workbook as _lw
    fpath = _resolve_file(spec["file"], base_dir, fallback_dir)
    if not fpath.exists():
        return []
    wb = _lw(str(fpath), read_only=True, data_only=True)
    sheet_name = spec.get("sheet", "")
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    skip = bool(spec.get("skip_header", True))
    rows: list[list] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i == 0 and skip:
            continue
        rows.append(list(row))
    wb.close()
    return rows


from typing import Optional
def _excel_write_run(data: dict) -> Optional[str]:
    """在录制时写入桌面 xlsx；失败返回错误信息。"""
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.utils import get_column_letter
    except ImportError:
        return "缺少 openpyxl：请执行 python3 rpa_manager.py deps-install B 或 pip install openpyxl / Missing openpyxl: run 'python3 rpa_manager.py deps-install B' or 'pip install openpyxl'"

    rel = (data.get("path") or data.get("value") or "").strip()
    sheet = (data.get("sheet") or "").strip()
    if not rel or not sheet:
        return "excel_write 需要 path（或 value）与 sheet / excel_write requires 'path' (or 'value') and 'sheet'"

    output_dir = Path.home() / "Desktop"  # recorder: final xlsx always to Desktop
    headers = data.get("headers") or []
    # Dynamic row sources (take precedence over static "rows")
    # rows_from_json / rows_from_excel 的中间文件在 tmp_dir，fallback 到 output_dir
    if data.get("rows_from_json"):
        rows = _excel_rows_from_json(data["rows_from_json"], _TASK_TMP_DIR, fallback_dir=output_dir)
    elif data.get("rows_from_excel"):
        rows = _excel_rows_from_excel(data["rows_from_excel"], _TASK_TMP_DIR, fallback_dir=output_dir)
    else:
        rows = data.get("rows") or []
    freeze = (data.get("freeze_panes") or "").strip() or None
    hidden_cols = data.get("hidden_columns") or []
    replace_sheet = bool(data.get("replace_sheet", True))

    path = output_dir / rel
    path.parent.mkdir(parents=True, exist_ok=True)

    if path.exists():
        wb = load_workbook(path)
    else:
        wb = Workbook()
        wb.remove(wb.active)

    if sheet in wb.sheetnames:
        if replace_sheet:
            del wb[sheet]
            ws = wb.create_sheet(sheet)
            ridx = 1
            for c, h in enumerate(headers, 1):
                ws.cell(row=ridx, column=c, value=h)
            if headers:
                ridx += 1
            for row in rows:
                for c, v in enumerate(row, 1):
                    ws.cell(row=ridx, column=c, value=v)
                ridx += 1
            if freeze:
                ws.freeze_panes = freeze
            for ci in hidden_cols:
                if isinstance(ci, int) and ci > 0:
                    ws.column_dimensions[get_column_letter(ci)].hidden = True
        else:
            ws = wb[sheet]
            nr = ws.max_row + 1
            for row in rows:
                for c, v in enumerate(row, 1):
                    ws.cell(row=nr, column=c, value=v)
                nr += 1
            if freeze:
                ws.freeze_panes = freeze
            for ci in hidden_cols:
                if isinstance(ci, int) and ci > 0:
                    ws.column_dimensions[get_column_letter(ci)].hidden = True
    else:
        ws = wb.create_sheet(sheet)
        ridx = 1
        for c, h in enumerate(headers, 1):
            ws.cell(row=ridx, column=c, value=h)
        if headers:
            ridx += 1
        for row in rows:
            for c, v in enumerate(row, 1):
                ws.cell(row=ridx, column=c, value=v)
            ridx += 1
        if freeze:
            ws.freeze_panes = freeze
        for ci in hidden_cols:
            if isinstance(ci, int) and ci > 0:
                ws.column_dimensions[get_column_letter(ci)].hidden = True

    wb.save(path)
    print(f"[recorder] excel_write → {path}", flush=True)
    return None


def _excel_write_codegen_lines(data: dict) -> list[str]:
    """与 _excel_write_run 同语义；生成 run() 内异步函数中的同步 openpyxl 代码。"""
    rel = (data.get("path") or data.get("value") or "").strip()
    sheet = (data.get("sheet") or "").strip()
    headers = data.get("headers") or []
    freeze = (data.get("freeze_panes") or "").strip() or None
    hidden_cols = data.get("hidden_columns") or []
    replace_sheet = bool(data.get("replace_sheet", True))

    h_repr = repr(headers)
    rel_repr = repr(rel)
    sheet_repr = repr(sheet)
    freeze_py = repr(freeze) if freeze else "None"
    hid_repr = repr(hidden_cols)
    rep_py = "True" if replace_sheet else "False"

    # --- Dynamic row-loading preamble ---
    row_preamble: list[str] = []
    rfj = data.get("rows_from_json")
    rfe = data.get("rows_from_excel")
    def _file_path_expr(file_str: str) -> str:
        """生成 Path 表达式：绝对路径直接 Path(...)，否则 tmp_dir（中间数据）。"""
        if Path(file_str).is_absolute():
            return f"Path({repr(file_str)})"
        return f'CONFIG["tmp_dir"] / {repr(file_str)}'

    if rfj:
        outer_key = rfj.get("outer_key", "")
        inner_key = rfj.get("inner_key", "")
        fields = rfj.get("fields") or []
        parent_fields = rfj.get("parent_fields") or []
        row_preamble += [
            f"_rfj_path = {_file_path_expr(rfj['file'])}",
            "import json as _json",
            "_rfj_data = _json.loads(_rfj_path.read_text(encoding='utf-8'))",
            "_rows = []",
        ]
        if inner_key:
            row_preamble += [
                f"for _outer in _rfj_data.get({repr(outer_key)}, []):",
                f"    _pvals = [_outer.get(_pf) for _pf in {repr(parent_fields)}]",
                f"    for _inner in _outer.get({repr(inner_key)}, []):",
                f"        _rows.append([_inner.get(_f) for _f in {repr(fields)}] + _pvals)",
            ]
        else:
            _src = f"_rfj_data.get({repr(outer_key)}, [])" if outer_key else "_rfj_data"
            row_preamble += [
                f"for _item in {_src}:",
                f"    _rows.append([_item.get(_f) for _f in {repr(fields)}]"
                + (f" + [_item.get(_pf) for _pf in {repr(parent_fields)}]" if parent_fields else "") + ")",
            ]
    elif rfe:
        src_sheet = repr(rfe.get("sheet", ""))
        skip = repr(bool(rfe.get("skip_header", True)))
        row_preamble += [
            f"_rfe_wb = load_workbook(str({_file_path_expr(rfe['file'])}), read_only=True, data_only=True)",
            f"_rfe_sn = {src_sheet}",
            "_rfe_ws = _rfe_wb[_rfe_sn] if _rfe_sn and _rfe_sn in _rfe_wb.sheetnames else _rfe_wb.active",
            "_rows = []",
            f"for _ri, _rrow in enumerate(_rfe_ws.iter_rows(values_only=True)):",
            f"    if _ri == 0 and {skip}: continue",
            "    _rows.append(list(_rrow))",
            "_rfe_wb.close()",
        ]
    else:
        rows = data.get("rows") or []
        row_preamble = [f"_rows = {repr(rows)}"]

    def _hide(indent: str) -> list[str]:
        return [
            indent + f"for _ci in {hid_repr}:",
            indent + "    if isinstance(_ci, int) and _ci > 0:",
            indent + "        _ws.column_dimensions[get_column_letter(_ci)].hidden = True",
        ]

    lines: list[str] = row_preamble + [
        "_xp = CONFIG[\"output_dir\"] / " + rel_repr,
        "_xp.parent.mkdir(parents=True, exist_ok=True)",
        "if _xp.exists():",
        "    _wb = load_workbook(_xp)",
        "else:",
        "    _wb = Workbook()",
        "    _wb.remove(_wb.active)",
        f"_sh = {sheet_repr}",
        f"_hdrs = {h_repr}",
        f"_replace = {rep_py}",
        "if _sh in _wb.sheetnames:",
        "    if _replace:",
        "        del _wb[_sh]",
        "        _ws = _wb.create_sheet(_sh)",
        "        _ridx = 1",
        "        for _c, _h in enumerate(_hdrs, 1):",
        "            _ws.cell(row=_ridx, column=_c, value=_h)",
        "        if _hdrs:",
        "            _ridx += 1",
        "        for _row in _rows:",
        "            for _c, _v in enumerate(_row, 1):",
        "                _ws.cell(row=_ridx, column=_c, value=_v)",
        "            _ridx += 1",
    ]
    if freeze:
        lines.append("        _ws.freeze_panes = " + freeze_py)
    lines.extend(_hide("        "))
    lines += [
        "    else:",
        "        _ws = _wb[_sh]",
        "        _nr = _ws.max_row + 1",
        "        for _row in _rows:",
        "            for _c, _v in enumerate(_row, 1):",
        "                _ws.cell(row=_nr, column=_c, value=_v)",
        "            _nr += 1",
    ]
    if freeze:
        lines.append("        _ws.freeze_panes = " + freeze_py)
    lines.extend(_hide("        "))
    lines += [
        "else:",
        "    _ws = _wb.create_sheet(_sh)",
        "    _ridx = 1",
        "    for _c, _h in enumerate(_hdrs, 1):",
        "        _ws.cell(row=_ridx, column=_c, value=_h)",
        "    if _hdrs:",
        "        _ridx += 1",
        "    for _row in _rows:",
        "        for _c, _v in enumerate(_row, 1):",
        "            _ws.cell(row=_ridx, column=_c, value=_v)",
        "        _ridx += 1",
    ]
    if freeze:
        lines.append("    _ws.freeze_panes = " + freeze_py)
    lines.extend(_hide("    "))
    lines += ["_wb.save(_xp)", 'print("excel_write →", _xp)']
    return lines


def _extract_parse_field_filenames(code: str) -> list[str]:
    """从代码里提取 _parse_field 调用中出现的所有文件名字符串（用于诊断）。

    匹配形如：
        _parse_field("xxx.txt", ...)
        _parse_field('xxx.txt', ...)
        _parse_field(CONFIG["output_dir"] / "xxx.txt", ...)
        _parse_field(Path('/tmp/xxx.txt'), ...)
    返回去重后的字符串列表（仅包含引号里的内容）。
    """
    import re as _re
    found = _re.findall(r'["\']([^"\']+\.[a-zA-Z]{2,6})["\']', code)
    # 过滤掉明显不是文件路径的（如 import 语句、print 格式串等）
    return list(dict.fromkeys(
        f for f in found
        if not f.startswith("utf") and "%" not in f and "\n" not in f
    ))


def _build_parse_field_example(file_registry: dict[str, list[str]]) -> str:
    """用本次 session 实际注册的文件和字段构建正确写法示例。"""
    if not file_registry:
        return (
            '      name = _parse_field(CONFIG["tmp_dir"] / "page_1.txt", "字段名")\n'
        )
    lines = []
    for fname, fields in list(file_registry.items())[:2]:  # 最多展示 2 个文件
        path_expr = f'CONFIG["tmp_dir"] / {repr(fname)}'
        for field in fields[:2]:  # 每文件最多 2 个字段示例
            var = field.replace(" ", "_").replace("/", "_")[:12]
            lines.append(f'      {var} = _parse_field({path_expr}, {repr(field)})')
    return "\n".join(lines) + "\n"


def _check_snippet_reads_extract_files(code: str) -> Optional[str]:
    """结构性卡口：验证 python_snippet 通过 _parse_field 读取本次 session 的提取文件。

    通用原则：
    - 仅当本次 session 存在 extract_text / extract_by_vision 步骤时生效
    - 代码必须调用 _parse_field(...)
    - 且引用的文件名至少包含一个已注册的提取文件

    豁免条件（任一满足则直接放行）：
    - 本次 session 无提取步骤（_EXTRACT_OUTPUT_FILES 为空）
    - 本次 session 已有 page.evaluate 成功步骤（_SESSION_HAS_PAGE_EVALUATE = True）：
      此时数据通道为 page.evaluate → JSON 文件，与 _parse_field 无关

    这不检测"写了什么坏代码"，而是验证"走了正确的数据通道"。
    兼容中英文字段名与文件名。
    Returns: 错误消息；None 表示通过。
    """
    if not _EXTRACT_OUTPUT_FILES:
        return None  # 本次 session 无提取步骤，不约束

    if _SESSION_HAS_PAGE_EVALUATE:
        return None  # page.evaluate 流：数据通道不经过 _parse_field，豁免

    # 1. 必须调用 _parse_field
    if "_parse_field" not in code:
        fields_hint = "; ".join(
            f"{fname}: {', '.join(fields)}"
            for fname, fields in _EXTRACT_FIELD_REGISTRY.items()
        )
        example = _build_parse_field_example(_EXTRACT_FIELD_REGISTRY)
        return (
            f"\n⛔  python_snippet 结构性卡口：\n"
            f"   本次 session 已用 extract_text / extract_by_vision 提取了数据到以下文件，\n"
            f"   但 python_snippet 没有调用 _parse_field() 读取任何提取文件。\n"
            f"\n"
            f"   已提取的文件与字段：\n"
            f"   {fields_hint}\n"
            f"\n"
            f"   ✅ 正确写法示例（使用本次 session 实际文件）：\n"
            f"{example}"
            f"\n"
            f"   _parse_field(filepath, field_name, index=0) 已注入到脚本，无需 import。\n"
            f"   / ⛔ Structural gate: no _parse_field() call found. "
            f"Use _parse_field to read extracted files.\n"
        )

    # 2. 引用的文件名至少匹配一个已注册提取文件
    referenced = [fname for fname in _EXTRACT_OUTPUT_FILES if fname in code]
    if not referenced:
        files_hint = ", ".join(f'"{f}"' for f in _EXTRACT_OUTPUT_FILES)
        # 诊断：代码里实际用了哪些文件名
        actual_used = _extract_parse_field_filenames(code)
        actual_hint = (
            f"   ⚠️  你的代码里出现的文件名：{', '.join(repr(f) for f in actual_used[:5])}\n"
            if actual_used else
            "   ⚠️  未从代码中识别到任何 .txt/.json 文件名引用。\n"
        )
        example = _build_parse_field_example(_EXTRACT_FIELD_REGISTRY)
        return (
            f"\n⛔  python_snippet 结构性卡口：\n"
            f"   代码中 _parse_field 的文件名参数未匹配本次 session 中任何已提取的文件。\n"
            f"\n"
            f"   本次 session 实际提取的文件：{files_hint}\n"
            f"{actual_hint}"
            f"\n"
            f"   ✅ 正确写法示例（使用本次 session 实际文件）：\n"
            f"{example}"
            f"\n"
            f"   请将 _parse_field 的第一个参数改为上述已提取文件名之一。\n"
            f"   / ⛔ _parse_field does not reference any registered extract file. "
            f"Expected one of: {files_hint}\n"
        )

    return None


class _PageEvaluateOnly:
    """python_snippet 沙箱中的限制性页面代理。

    仅开放 page.evaluate(JS)，用于列表行型数据一次性行对齐提取。
    其他所有页面方法（locator / click / fill / goto 等）仍被阻断——
    交互操作必须使用专用 action，不应在 python_snippet 里内联。
    """

    def __init__(self, real_page):
        object.__setattr__(self, "_real_page", real_page)

    def evaluate(self, *args, **kwargs):
        """Proxy to real Playwright page.evaluate() — returns an awaitable coroutine."""
        real = object.__getattribute__(self, "_real_page")
        return real.evaluate(*args, **kwargs)

    def __getattr__(self, name: str):
        raise RuntimeError(
            f"\n❌  python_snippet: page.{name}() 不可用。\n"
            f"\n"
            f"python_snippet 里只允许 page.evaluate(JS字符串) 用于列表行型数据提取。\n"
            f"其他交互（点击/填写/导航）请用专用 action：click / fill / goto 等。\n"
            f"\n"
            f"/ python_snippet: only page.evaluate(JS) is allowed for list-row extraction.\n"
            f"  Use dedicated actions (click / fill / goto) for browser interaction.\n"
        )

    def __bool__(self):
        return True


class _MockPage:
    """python_snippet 沙箱降级用：无浏览器上下文时的完全阻断占位对象。

    正常录制流程应使用 _PageEvaluateOnly（传入真实 page）。
    仅当 page 不可用时（非浏览器任务）回退到此对象。
    """

    _MSG = (
        "\n"
        "❌  python_snippet 无法访问页面（当前任务无浏览器上下文）。\n"
        "\n"
        "如需提取网页数据，请在 python_snippet 之前先发送：\n"
        "  - extract_text action（SSR 页面单字段提取）\n"
        "  - extract_by_vision action（视觉提取，适合 SPA 或列表行型）\n"
        "\n"
        "python_snippet 只能做：文件读写 / 数据解析 / datetime / openpyxl / docx。\n"
    )

    def __getattr__(self, name: str):
        raise RuntimeError(self._MSG)

    def __bool__(self):
        return False


async def _python_snippet_run(code: str, page) -> Optional[str]:
    """在录制时 async 执行 python_snippet 代码，验证依赖和逻辑正确性。

    代码被包裹进 async def __snippet__() 中执行，因此：
      - await page.evaluate(JS)  ✅ 支持（列表行型提取）
      - 普通同步文件读写           ✅ 支持
      - page.locator / page.click 等  ❌ 仍被 _PageEvaluateOnly 阻断

    返回错误字符串；None 表示成功。
    """
    import traceback as _tb

    # Wrap original code in async function so `await` is valid at any level
    indented_lines = []
    for line in code.splitlines():
        indented_lines.append(("    " + line) if line.strip() else "")
    async_code = "async def __snippet__():\n" + "\n".join(indented_lines) + "\n"

    ns: dict = {
        "Path": Path,
        "CONFIG": {
            "output_dir": Path.home() / "Desktop",
            "tmp_dir":    _TASK_TMP_DIR,
            "task_name":  "preview",
        },
        "json": __import__("json"),
        "os": __import__("os"),
        "re": __import__("re"),
        "datetime": __import__("datetime"),
        "page": _PageEvaluateOnly(page) if page is not None else _MockPage(),
        "_parse_field": _parse_field,
    }

    # openpyxl
    _missing_deps: list[str] = []
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.utils import get_column_letter
        ns.update({"Workbook": Workbook, "load_workbook": load_workbook,
                   "get_column_letter": get_column_letter})
    except ImportError:
        _missing_deps.append("openpyxl")

    # python-docx
    try:
        from docx import Document
        ns["Document"] = Document
    except ImportError:
        _missing_deps.append("python-docx")

    # Pre-check: if code references openpyxl/docx symbols but dep is missing, fail fast
    needs_openpyxl = any(sym in code for sym in ("load_workbook", "Workbook", "openpyxl", "get_column_letter"))
    needs_docx = any(sym in code for sym in ("Document", "from docx", "python_docx"))
    missing_required = []
    if needs_openpyxl and "openpyxl" in _missing_deps:
        missing_required.append("openpyxl")
    if needs_docx and "python-docx" in _missing_deps:
        missing_required.append("python-docx")
    if missing_required:
        pkgs = " ".join(missing_required)
        cap = "B" if "openpyxl" in missing_required else "C"
        return (f"python_snippet 缺少依赖：{pkgs}。请先执行：python3 rpa_manager.py deps-install {cap}"
                f" / python_snippet missing deps: {pkgs}. Run: python3 rpa_manager.py deps-install {cap}")

    # Compile (syntax check on wrapped code; adjust reported line by -1 for wrapper)
    try:
        compiled = compile(async_code, "<python_snippet>", "exec")
    except SyntaxError as e:
        lineno = max(1, (e.lineno or 1) - 1)
        return f"python_snippet 语法错误 / syntax error: line {lineno}: {e.msg}"

    # Execute: define __snippet__ then await it
    try:
        exec(compiled, ns)          # noqa: S102  — defines __snippet__ in ns
        await ns["__snippet__"]()   # run the async snippet
        print("[recorder] python_snippet 验证通过 / validation passed ✓", flush=True)
        return None
    except ImportError as e:
        mod = str(e).split("'")[1] if "'" in str(e) else str(e)
        return (f"python_snippet ImportError：{e}。请确认 {mod} 已安装（deps-install 或 pip install {mod}）"
                f" / ImportError: {e}. Make sure '{mod}' is installed (deps-install or pip install {mod}).")
    except FileNotFoundError as e:
        return (f"python_snippet 文件未找到：{e}。请确认前序步骤（api_call / excel_write）已成功执行并在桌面生成了该文件"
                f" / FileNotFoundError: {e}. Make sure the preceding api_call / excel_write step ran successfully and the file exists on the Desktop.")
    except Exception as e:
        tb = _tb.format_exc(limit=5)
        return f"python_snippet 执行失败 / execution failed: {type(e).__name__}: {e}\n{tb}"


def _word_write_run(data: dict) -> Optional[str]:
    try:
        from docx import Document
    except ImportError:
        return "缺少 python-docx：请执行 python3 rpa_manager.py deps-install C 或 pip install python-docx / Missing python-docx: run 'python3 rpa_manager.py deps-install C' or 'pip install python-docx'"

    # Accept "path", "target" (common AI alias), or legacy fallback "value" for the file path.
    rel = (data.get("path") or data.get("target") or data.get("value") or "").strip()
    if not rel:
        return "word_write 需要 path（或 value）/ word_write requires 'path' (or 'value')"

    paragraphs = data.get("paragraphs") or []
    if not isinstance(paragraphs, list):
        return "paragraphs 须为字符串数组 / 'paragraphs' must be an array of strings"

    table_def = data.get("table")  # optional: {"headers": [...], "rows": [[...]]}

    mode = (data.get("mode") or "new").lower()
    # 路径：含 ~ 或 / 前缀视为绝对路径展开，否则落到 ~/Desktop/
    if rel.startswith("~") or rel.startswith("/"):
        path = Path(rel).expanduser()
    else:
        path = Path.home() / "Desktop" / rel
    path.parent.mkdir(parents=True, exist_ok=True)

    try:
        if mode == "append" and path.exists():
            doc = Document(str(path))
        else:
            doc = Document()
        for p in paragraphs:
            doc.add_paragraph(str(p))
        if table_def and isinstance(table_def, dict):
            headers = table_def.get("headers") or []
            rows_from_json = table_def.get("rows_from_json")
            if rows_from_json and isinstance(rows_from_json, dict):
                import json as _j
                rjf = rows_from_json.get("file", "")
                if rjf.startswith("~") or rjf.startswith("/"):
                    rjf_path = Path(rjf).expanduser()
                else:
                    # 中间数据优先在 tmp_dir 查找，不存在则 fallback 到 Desktop
                    rjf_path = _TASK_TMP_DIR / rjf
                    if not rjf_path.exists():
                        rjf_path = Path.home() / "Desktop" / rjf
                if not rjf_path.exists():
                    return f"word_write rows_from_json: 文件不存在 {rjf_path}"
                rows = _j.loads(rjf_path.read_text(encoding="utf-8"))
                if not isinstance(rows, list):
                    return f"word_write rows_from_json: 文件内容须为 JSON 数组 {rjf_path}"
            else:
                rows = table_def.get("rows") or []
            col_count = len(headers) or (len(rows[0]) if rows else 0)
            if col_count:
                tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
                tbl.style = "Table Grid"
                hdr_cells = tbl.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h)
                for r_idx, row in enumerate(rows):
                    cells = tbl.rows[r_idx + 1].cells
                    for c_idx, val in enumerate(row):
                        cells[c_idx].text = str(val)
        doc.save(str(path))
        print(f"[recorder] word_write → {path}", flush=True)
    except Exception as exc:
        return str(exc)
    return None


def _expand_para_placeholders(para: str) -> str:
    """将段落文本里的 {{now:fmt}} 替换为运行时 datetime 表达式字符串。

    规则：
      {{now}}           → datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
      {{now:%m月%d日}}  → datetime.datetime.now().strftime('%m月%d日')
    返回的是一段 Python f-string 表达式，供 codegen 嵌入生成代码。
    如果段落里没有 {{now...}}，直接返回 repr(para)（普通字符串字面量）。
    """
    import re as _re
    _NOW_PAT = _re.compile(r"\{\{now(?::([^}]*))?\}\}")
    if not _NOW_PAT.search(para):
        return repr(para)
    # 把 para 拆成普通文本片段和 {{now:...}} 片段，拼成 f-string
    parts = []
    last = 0
    for m in _NOW_PAT.finditer(para):
        if m.start() > last:
            parts.append(repr(para[last:m.start()]))
        fmt = m.group(1) or "%Y-%m-%d %H:%M:%S"
        parts.append(f'__import__("datetime").datetime.now().strftime({repr(fmt)})')
        last = m.end()
    if last < len(para):
        parts.append(repr(para[last:]))
    return "(" + " + ".join(parts) + ")"


def _word_write_codegen_lines(data: dict) -> list[str]:
    # Accept "path", "target" (common AI alias), or legacy fallback "value".
    rel = (data.get("path") or data.get("target") or data.get("value") or "").strip()
    paragraphs = data.get("paragraphs") or []
    table_def = data.get("table")
    mode = (data.get("mode") or "new").lower()
    mode_repr = repr(mode)

    # 路径生成：若含 ~ 或 / 前缀视为绝对路径，展开后直接使用；否则拼 output_dir
    if rel.startswith("~") or rel.startswith("/"):
        wp_line = f"_wp = Path({repr(rel)}).expanduser()"
    else:
        wp_line = "_wp = CONFIG[\"output_dir\"] / " + repr(rel)

    # paragraphs 支持 {{now:fmt}} 动态占位符，其余字符串仍序列化为字面量
    para_exprs = [_expand_para_placeholders(str(p)) for p in paragraphs]
    para_list_code = "[" + ", ".join(para_exprs) + "]"

    lines = [
        wp_line,
        "_wp.parent.mkdir(parents=True, exist_ok=True)",
        f"_wparas = {para_list_code}",
        f"_wmode = {mode_repr}",
        'if _wmode == "append" and _wp.exists():',
        "    _doc = Document(str(_wp))",
        "else:",
        "    _doc = Document()",
        "for _p in _wparas:",
        "    _doc.add_paragraph(str(_p))",
    ]

    if table_def and isinstance(table_def, dict):
        headers = table_def.get("headers") or []
        rows = table_def.get("rows") or []
        rows_from_json = table_def.get("rows_from_json")  # 动态来源

        if rows_from_json and isinstance(rows_from_json, dict):
            # 动态模式：从中间 JSON 文件读取 rows（与 excel_write 保持一致）
            rjf = rows_from_json.get("file", "")
            rjf_repr = repr(rjf)
            # 路径生成与 _word_write_run 保持一致：~ 或 / 开头视为绝对路径展开，
            # 否则从 tmp_dir 读取（中间数据），不存在时 fallback 到 output_dir
            if rjf.startswith("~") or rjf.startswith("/"):
                src_line = f"_wtbl_src = Path({rjf_repr}).expanduser()"
            else:
                src_line = (
                    f"_wtbl_src = CONFIG[\"tmp_dir\"] / {rjf_repr}; "
                    f"_wtbl_src = _wtbl_src if _wtbl_src.exists() else CONFIG[\"output_dir\"] / {rjf_repr}"
                )
            lines += [
                f"_wtbl_headers = {repr(headers)}",
                "import json as _wtbl_json",
                src_line,
                "_wtbl_rows = _wtbl_json.loads(_wtbl_src.read_text(encoding='utf-8')) if _wtbl_src.exists() else []",
            ]
        else:
            # 静态模式：rows 直接写入（仅适用于真正静态的模板数据）
            if rows:
                lines += [
                    "# ⚠ rows 为录制时填入的静态数据，仅适用于真正不变的模板行",
                    "# 如果此数据来自网页提取，应改用 python_snippet + _parse_field 动态构建",
                ]
            lines += [
                f"_wtbl_headers = {repr(headers)}",
                f"_wtbl_rows = {repr(rows)}",
            ]

        lines += [
            "_wtbl_cols = len(_wtbl_headers) or (len(_wtbl_rows[0]) if _wtbl_rows else 0)",
            "if _wtbl_cols:",
            "    _wtbl = _doc.add_table(rows=1 + len(_wtbl_rows), cols=_wtbl_cols)",
            '    _wtbl.style = "Table Grid"',
            "    for _ci, _h in enumerate(_wtbl_headers):",
            "        _wtbl.rows[0].cells[_ci].text = str(_h)",
            "    for _ri, _row in enumerate(_wtbl_rows):",
            "        for _ci, _v in enumerate(_row):",
            "            _wtbl.rows[_ri + 1].cells[_ci].text = str(_v)",
        ]

    lines += [
        "_doc.save(str(_wp))",
        'print("word_write →", _wp)',
    ]
    return lines


def _write_kv_field(out_path: Path, field_name: str, values: list[str], first_write: bool) -> None:
    """将 extract_text 提取结果写入 kv 格式临时文件。

    格式规则（兼容中英文字段名与字段值）：
      单值：  field_name: value
      多值：  field_name.0: value0
              field_name.1: value1
              ...
    追加写入时直接续行；调用方负责 first_write 标志。
    """
    if not values:
        return  # 0 条时不写任何内容（read 时 raise RuntimeError 提示选择器问题）
    lines: list[str] = []
    if len(values) == 1:
        lines.append(f"{field_name}: {values[0]}")
    else:
        for i, v in enumerate(values):
            lines.append(f"{field_name}.{i}: {v}")
    blob = "\n".join(lines) + "\n"
    if first_write:
        out_path.write_text(blob, encoding="utf-8")
    else:
        with out_path.open("a", encoding="utf-8") as f:
            f.write(blob)


def _parse_field(filepath, field_name: str, index: int = 0):
    """从 extract_text 输出的 kv 文件中读取指定字段值。

    兼容中英文字段名与值（UTF-8）。支持带 .N 索引后缀的多值字段。

    Args:
        filepath:   文件路径（str 或 Path）
        field_name: 字段名，与 extract_text 的 field 参数一致
        index:      0 = 第一条（默认）；-1 = 最后一条；None = 返回全部列表

    Raises:
        RuntimeError: 文件不存在，或字段在文件中未找到
    """
    path = Path(filepath) if not isinstance(filepath, Path) else filepath
    if not path.exists():
        raise RuntimeError(
            f"提取文件不存在 / Extract file not found: {path}\n"
            f"请确认 extract_text 步骤已成功执行并写入该文件 / "
            f"Make sure the extract_text step ran successfully and wrote this file."
        )
    matches: list[str] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if ":" not in line:
            continue
        k, _, v = line.partition(":")
        k_base = k.strip()
        # 去掉 .N 索引后缀：field.0 → field
        if "." in k_base and k_base.rsplit(".", 1)[-1].isdigit():
            k_base = k_base.rsplit(".", 1)[0]
        if k_base == field_name:
            matches.append(v.strip())
    if not matches:
        raise RuntimeError(
            f"字段 '{field_name}' 在 {path} 中未找到 / "
            f"Field '{field_name}' not found in {path}.\n"
            f"请检查 extract_text 的 field 参数是否与此处一致 / "
            f"Check that the extract_text 'field' param matches this name."
        )
    if index is None:
        return matches
    try:
        return matches[index]
    except IndexError:
        return matches[-1]


# ── DOM snapshot ─────────────────────────────────────────────────────────────

async def _snapshot(page) -> dict:
    """Return DOM intelligence for LLM selector decisions.

    Always returns three layers:
      items       – interactive + heading elements (navigation / interaction layer)
      sections    – named content regions (structure layer)
      data_groups – auto-detected repeating data containers with sampled field
                    selectors (data layer).  The LLM uses these to write correct
                    page.evaluate() JS without guessing.  Works on any website;
                    detection is pure DOM structural analysis, no site-specific
                    knowledge required.

    Selector priority for items:
      1. Own  #id / [data-testid] / [aria-label] / tag[name]
      2. Ancestor walk (max 4 levels) — produces  [data-testid="X"] tag
      3. :nth-of-type fallback inside nearest sectioning parent
    """
    try:
        result = await page.evaluate("""() => {
            // ── helpers ──────────────────────────────────────────────────────
            function ownSel(el) {
                if (el.id) return '#' + el.id;
                const tid  = el.getAttribute('data-testid');
                if (tid)  return `[data-testid="${tid}"]`;
                const aria = el.getAttribute('aria-label');
                if (aria) return `[aria-label="${aria}"]`;
                const name = el.getAttribute('name');
                if (name) return `${el.tagName.toLowerCase()}[name="${name}"]`;
                return null;
            }

            function ancestorSel(el) {
                let cur = el.parentElement;
                const midTags = [];
                for (let d = 0; d < 4 && cur; d++, cur = cur.parentElement) {
                    const s = ownSel(cur);
                    if (s) {
                        const mid = midTags.slice().reverse().join(' ');
                        return mid ? `${s} ${mid} ${el.tagName.toLowerCase()}`
                                   : `${s} ${el.tagName.toLowerCase()}`;
                    }
                    midTags.push(cur.tagName.toLowerCase());
                }
                return null;
            }

            function nthSel(el) {
                const parent = el.parentElement;
                if (!parent) return null;
                const siblings = Array.from(parent.children)
                    .filter(c => c.tagName === el.tagName);
                const idx = siblings.indexOf(el) + 1;
                const ps = ownSel(parent);
                if (ps) return `${ps} > ${el.tagName.toLowerCase()}:nth-of-type(${idx})`;
                return null;
            }

            function isVisible(el) {
                const r = el.getBoundingClientRect();
                return r.width > 0 && r.height > 0;
            }

            // ── Layer 1: interactive + heading elements ───────────────────
            const TAGS = [
                'input', 'button', 'select', 'textarea', 'a[href]',
                '[role="button"]', '[role="link"]', '[role="searchbox"]',
                '[role="tab"]', 'h1', 'h2', 'h3', 'li'
            ].join(',');

            const items = Array.from(document.querySelectorAll(TAGS))
                .filter(isVisible)
                .slice(0, 100)
                .map(el => {
                    const tag  = el.tagName.toLowerCase();
                    const ph   = el.getAttribute('placeholder') || null;
                    const text = (el.textContent || '').replace(/\\s+/g, ' ').trim().slice(0, 70);
                    const sel  = ownSel(el) || ancestorSel(el) || nthSel(el);
                    return { tag, sel: sel || null, ph, text: text || null };
                }).filter(e => e.sel || e.text);

            // ── Layer 2: named content sections ──────────────────────────
            const SECTION_TAGS = [
                'section', 'article', '[data-testid]', '[id]'
            ].join(',');
            const sections = Array.from(document.querySelectorAll(SECTION_TAGS))
                .filter(el => {
                    const r = el.getBoundingClientRect();
                    return r.width > 100 && r.height > 50;
                })
                .slice(0, 20)
                .map(el => {
                    const s   = ownSel(el);
                    const h   = el.querySelector('h1,h2,h3');
                    const heading = h ? (h.textContent||'').replace(/\\s+/g,' ').trim().slice(0,50) : null;
                    return s ? { sel: s, heading } : null;
                })
                .filter(Boolean)
                .filter((v, i, a) => a.findIndex(x => x.sel === v.sel) === i);

            // ── Layer 3: repeating data containers ───────────────────────
            //
            // Universal detection — works on any website, no site-specific knowledge.
            // Any website's list data is structurally: sibling elements sharing the
            // same pattern under a common parent.  Two complementary strategies:
            //
            //  A) Semantic-attribute groups — elements sharing the same data-* /
            //     itemtype / role attribute value, repeated 3+ times.
            //     Covers: Amazon (data-component-type), React/Vue (data-cy,
            //     data-item-id), schema.org (itemtype), ARIA lists (role="listitem").
            //
            //  B) Structural-similarity groups — siblings under the same parent
            //     with the same tagName + first CSS class, repeated 4+ times.
            //     Covers: class-based lists (ul > li.card, div.grid > div.item).
            //
            // First element of each group is sampled to extract visible text-bearing
            // and link descendants with relative selectors.
            // The LLM uses sample_fields to write correct page.evaluate() JS directly
            // without guessing or needing a separate dom_inspect step.

            // ── helper: build a minimal CSS selector string for an element ──
            function elSel(el) {
                if (el.id) return '#' + CSS.escape(el.id);
                // semantic data attributes (most specific first)
                for (const attr of ['data-component-type','data-testid','data-cy',
                                     'data-type','data-item-id','data-asin','itemtype']) {
                    const v = el.getAttribute(attr);
                    if (v) return `${el.tagName.toLowerCase()}[${attr}="${v}"]`;
                }
                // class-based fallback
                const cls = Array.from(el.classList).slice(0,2).join('.');
                return cls ? `${el.tagName.toLowerCase()}.${cls}` : el.tagName.toLowerCase();
            }

            // ── helper: selector relative to a container ancestor ──────────
            function relSel(container, child) {
                const parts = [];
                let cur = child;
                while (cur && cur !== container) {
                    // prefer distinguishing attr/class over plain tag
                    let part = null;
                    for (const attr of ['data-testid','data-cy','aria-label','name','itemprop']) {
                        const v = cur.getAttribute(attr);
                        if (v) { part = `[${attr}="${v}"]`; break; }
                    }
                    if (!part) {
                        const cls = cur.classList[0];
                        part = cls ? `${cur.tagName.toLowerCase()}.${CSS.escape(cls)}`
                                   : cur.tagName.toLowerCase();
                    }
                    parts.unshift(part);
                    cur = cur.parentElement;
                }
                return parts.join(' > ');
            }

            // ── helper: sample visible text/link fields from a container ──
            function sampleFields(container) {
                const fields = [];
                const seen = new Set();
                // Walk text-bearing and link children (ordered by DOM position)
                const candidates = Array.from(
                    container.querySelectorAll('h1,h2,h3,h4,h5,h6,span,p,a,img,[itemprop]')
                );
                for (const child of candidates) {
                    if (!isVisible(child)) continue;
                    // skip deeply nested duplicates already covered by ancestor
                    const rSel = relSel(container, child);
                    if (!rSel || seen.has(rSel)) continue;
                    seen.add(rSel);

                    const tag  = child.tagName.toLowerCase();
                    const text = (child.textContent||'').replace(/\\s+/g,' ').trim().slice(0,80);
                    // Use .href property (absolute URL) for <a> elements; fall back to attribute
                    const rawHref = child.href || child.getAttribute('href') || '';
                    // Skip javascript: voids and empty anchors — they are UI chrome, not data links
                    const href = (rawHref && !rawHref.startsWith('javascript:')) ? rawHref : '';
                    const alt  = child.getAttribute('alt');
                    const itemprop = child.getAttribute('itemprop');

                    // only include nodes that carry meaningful content
                    if (!text && !href && !alt) continue;
                    if (text.length < 2 && !href) continue;

                    const field = { sel: rSel, tag };
                    if (text)      field.text     = text;
                    if (href)      field.href      = href.slice(0, 200);   // full absolute URL
                    if (alt)       field.alt       = alt.slice(0, 60);
                    if (itemprop)  field.itemprop  = itemprop;
                    fields.push(field);
                    if (fields.length >= 12) break;
                }
                return fields;
            }

            // ── Strategy A: semantic-attribute groups ─────────────────────
            const SEMANTIC_ATTRS = [
                'data-component-type', 'data-asin', 'data-cy', 'data-type',
                'data-item-id', 'data-item', 'data-row', 'data-index',
                'itemtype', 'role'
            ];
            const semGroups = {};   // key → [el, ...]
            for (const attr of SEMANTIC_ATTRS) {
                document.querySelectorAll(`[${attr}]`).forEach(el => {
                    const val = el.getAttribute(attr);
                    // skip generic roles that are not list-item roles
                    if (attr === 'role' && !['listitem','row','gridcell','article','option'].includes(val)) return;
                    const key = `${el.tagName.toLowerCase()}[${attr}="${val}"]`;
                    if (!semGroups[key]) semGroups[key] = [];
                    semGroups[key].push(el);
                });
            }

            // ── Strategy B: structural-similarity groups ──────────────────
            // Group same-parent children by (tagName + firstClass) combos
            const structGroups = {};
            document.querySelectorAll('ul > *,ol > *,tbody > tr,[class] > [class]').forEach(el => {
                const parent = el.parentElement;
                if (!parent) return;
                const firstCls = el.classList[0] || '';
                if (!firstCls && el.tagName === 'LI') {
                    // plain <li> without class: group by parent
                    const key = `parent:${elSel(parent)} > li`;
                    if (!structGroups[key]) structGroups[key] = [];
                    if (!structGroups[key].includes(el)) structGroups[key].push(el);
                    return;
                }
                if (!firstCls) return;
                const key = `parent:${elSel(parent)} > ${el.tagName.toLowerCase()}.${CSS.escape(firstCls)}`;
                if (!structGroups[key]) structGroups[key] = [];
                if (!structGroups[key].includes(el)) structGroups[key].push(el);
            });

            // ── data richness scorer ──────────────────────────────────────
            // Ranks groups by how useful they are as data containers.
            // Goal: surface the actual "data row" (e.g. product card, article card)
            // rather than sub-components (title slot, add-to-cart button, rating icon).
            // Known non-data component type patterns: tracking, analytics, ads, chrome.
            // Groups whose container_sel key contains these strings are penalised heavily.
            const TRACKER_PATTERNS = [
                'impression', 'logger', 'counter', 'tracker', 'analytics',
                'beacon', 'pixel', 'sponsor', 'ad-slot', 'carousel-slide',
            ];

            function scoreGroup(key, g, fields) {
                let score = 0;
                const count = g.els.length;

                // Semantic attr groups (Strategy A) are far more likely to be
                // meaningful data containers than class-based structural groups (B).
                if (key.startsWith('A:')) score += 5;

                // Tracker/analytics component penalty — these are never data containers.
                const lowerKey = key.toLowerCase();
                if (TRACKER_PATTERNS.some(p => lowerKey.includes(p))) score -= 8;

                // Count scoring: prefer lists in the 20-100 range.
                // Very small (3-10) might be anything; very large (200+) are sub-components.
                if      (count >= 20  && count <= 100) score += 5;  // ideal data list
                else if (count >= 10  && count <  20)  score += 3;
                else if (count >= 101 && count <= 200)  score += 1;
                else if (count > 200)                   score -= 3;  // sub-component
                else                                    score += 1;  // count 3-9

                // Field diversity: links + varied text = rich data container
                const hasHref    = fields.some(f => f.href);
                const hasText    = fields.some(f => f.text && f.text.length > 5);
                const textValues = fields.map(f => (f.text || '').trim()).filter(Boolean);
                const uniqueTexts = new Set(textValues).size;
                // UI widget: all fields carry the same text (button label, nav item…)
                const isUIWidget = textValues.length >= 2 && uniqueTexts <= 1;

                if (hasHref)            score += 3;   // URLs are extractable data
                if (hasText)            score += 1;
                if (fields.length >= 3) score += 2;   // multiple distinct fields
                if (fields.length >= 6) score += 1;   // extra rich
                if (isUIWidget)         score -= 6;   // heavy penalty for button/nav groups

                return score;
            }

            // ── merge & score groups ──────────────────────────────────────
            const allGroups = {};
            Object.entries(semGroups).forEach(([k, els])    => { allGroups['A:' + k] = { els, strategy: 'semantic' }; });
            Object.entries(structGroups).forEach(([k, els]) => { allGroups['B:' + k] = { els, strategy: 'structural' }; });

            const MIN_COUNT = 3;

            // First pass: sample fields and compute score for every candidate group
            const scored = Object.entries(allGroups)
                .filter(([, g]) => g.els.length >= MIN_COUNT)
                .map(([key, g]) => {
                    const firstEl = g.els.find(isVisible) || g.els[0];
                    const fields  = sampleFields(firstEl);
                    const score   = scoreGroup(key, g, fields);
                    return { key, g, fields, score, _elSet: new Set(g.els) };
                })
                .filter(c => c.fields.length > 0)         // must have at least one field
                .sort((a, b) => b.score - a.score);        // best containers first

            // Second pass: deduplicate — skip if all elements already in a higher-scored group
            const data_groups = scored
                .reduce((acc, c) => {
                    const alreadyCovered = acc.some(existing =>
                        c.g.els.every(el => existing._elSet.has(el))
                    );
                    if (!alreadyCovered) acc.push(c);
                    return acc;
                }, [])
                .slice(0, 10)                              // return up to 10 groups
                .map(({ key, g, fields }) => ({
                    container_sel: elSel(g.els.find(isVisible) || g.els[0]),
                    count:         g.els.length,
                    strategy:      g.strategy,
                    sample_fields: fields,
                }));

            return { items, sections, data_groups };
        }""")
        return result if isinstance(result, dict) else {"items": result or [], "sections": [], "data_groups": []}
    except Exception:
        return {"items": [], "sections": [], "data_groups": []}


# ── Action executor ──────────────────────────────────────────────────────────

async def _do_action(page, data: dict, step_n: int, shots_dir: Path) -> dict:
    # seq=0 means a brand-new recording session started; reset all session-level globals
    # to prevent stale state from a previous run bleeding into this session.
    if data.get("seq", -1) == 0:
        _reset_extract_output_tracking()

    action  = data.get("action", "")
    target  = data.get("target") or data.get("url", "")   # "url" is a common LLM alias for "target"
    value   = data.get("value", "")
    context = data.get("context") or f"步骤 {step_n}"

    code_block        = None
    error             = None
    inspect_children  = None  # dom_inspect: passed through to result JSON for rpa_manager
    extract_summary   = None  # extract_text: structured summary (no raw values) for AI

    # No-browser mode: browser-specific actions are not available.
    # 无浏览器模式：浏览器操作不可用，返回明确错误提示。
    _BROWSER_ACTIONS = {"goto", "fill", "press", "click", "select_option", "extract_text",
                        "extract_by_vision", "wait", "scroll", "scroll_to", "snapshot", "dom_inspect"}
    if page is None and action in _BROWSER_ACTIONS:
        result: dict = {
            "success":    False,
            "error":      (
                f"此任务能力码不含浏览器，不支持 {action!r} 操作。"
                f" / This capability does not include a browser; {action!r} is not available."
            ),
            "code_block": None,
            "screenshot": None,
            "url":        "",
            "snapshot":   [],
            "sections":   [],
        }
        return result

    try:
        if action == "goto":
            await page.goto(target, wait_until="domcontentloaded")
            await page.wait_for_timeout(1500)   # SPA initial render
            code_block = _step_code(step_n, context, [
                f'await page.goto({repr(target)}, wait_until="domcontentloaded")',
                'await page.wait_for_timeout(CONFIG["spa_settle_ms"])',
            ])

        elif action == "fill":
            loc = page.locator(target).first
            await loc.wait_for(state="visible", timeout=20_000)
            await loc.fill(value)
            code_block = _step_code(step_n, context, [
                f'await page.locator({repr(target)}).first.fill({repr(value)})',
            ])

        elif action == "press":
            key = target or "Enter"
            await page.keyboard.press(key)
            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(800)   # let SPA finish routing
            code_block = _step_code(step_n, context, [
                f'await page.keyboard.press({repr(key)})',
                'await page.wait_for_load_state("domcontentloaded")',
                'await page.wait_for_timeout(800)',
            ])

        elif action == "click":
            loc = page.locator(target).first
            await loc.wait_for(state="visible", timeout=20_000)
            await loc.click()
            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(800)   # let SPA finish routing
            code_block = _step_code(step_n, context, [
                f'await page.locator({repr(target)}).first.click()',
                'await page.wait_for_load_state("domcontentloaded")',
                'await page.wait_for_timeout(800)',
            ])

        elif action == "select_option":
            # Native <select>: target = CSS, value = option value / label / index (see select_by)
            # Playwright fill() does NOT set <select>; use select_option (e.g. Sauce Demo hilo = price high→low)
            loc = page.locator(target).first
            await loc.wait_for(state="visible", timeout=20_000)
            how = (data.get("select_by") or "value").lower().strip()
            if how == "label":
                await loc.select_option(label=value)
                sel_line = f'await page.locator({repr(target)}).first.select_option(label={repr(value)})'
            elif how == "index":
                idx = int(value)
                await loc.select_option(index=idx)
                sel_line = f'await page.locator({repr(target)}).first.select_option(index={idx})'
            else:
                await loc.select_option(value)
                sel_line = f'await page.locator({repr(target)}).first.select_option({repr(value)})'
            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(800)
            code_block = _step_code(step_n, context, [
                sel_line,
                'await page.wait_for_load_state("domcontentloaded")',
                'await page.wait_for_timeout(800)',
            ])

        elif action == "extract_text":
            filename = value or "output.txt"
            limit    = int(data.get("limit", 0)) or 0
            extract_summary = None
            code_block = None

            try:
                _xh = (urllib.parse.urlparse(page.url or "").hostname or "").lower()
            except Exception:
                _xh = ""
            _force_dom = bool(data.get("force_extract_text"))

            if _hostname_on_heavy_spa_list(_xh) and not _force_dom:
                error = (
                    f"当前页 hostname「{_xh}」命中内置重型 SPA 列表；从该站提取字段必须使用 **extract_by_vision**，"
                    f"不要使用 extract_text（哈希 class 与技能拆解阶段「默认视觉」一致，避免无效录制）。"
                    f"\n确需 DOM 提取时可在 JSON 中加 \"force_extract_text\": true（不推荐）。"
                    f"\n / Host is on the heavy-SPA list; use extract_by_vision for field extraction, not extract_text."
                )
                print(f"[recorder] extract_text 已拒绝（重型 SPA）hostname={_xh!r}", flush=True)
            else:
                # 与视觉提取相同：先等 SPA 主内容（避免骨架屏上 querySelector 得到空节点）
                _er_ms = int(data.get("extract_ready_timeout_ms") or 30_000)
                await _wait_spa_ready_for_vision(page, "", timeout_ms=_er_ms)

                # Wait for page to settle (SPA re-renders can cause locator.all() race)
                try:
                    await page.wait_for_load_state("domcontentloaded", timeout=5_000)
                except Exception:
                    pass

                # Single atomic JS call — immune to mid-flight page re-renders
                limit_n = limit or 9999
                texts = await page.evaluate(_EXTRACT_JS_MIN, [target, limit_n])

                # field / field_name = short label for kv key; else fall back to context
                field_label = (
                    (data.get("field") or data.get("field_name") or context or f"步骤 {step_n}")
                    .strip()
                    or "extract"
                )
                first_for_name = filename not in _EXTRACT_OUTPUT_FILES
                _EXTRACT_OUTPUT_FILES.add(filename)

                # 中间提取文件统一写入 tmp_dir（隔离不同任务）
                out = _TASK_TMP_DIR / filename
                # 写 kv 格式（不暴露完整值给 AI 上下文）
                _write_kv_field(out, field_label, texts, first_write=first_for_name)
                # 注册字段名，供 python_snippet 结构性卡口使用
                _EXTRACT_FIELD_REGISTRY.setdefault(filename, []).append(field_label)

                if texts:
                    print(
                        f"[recorder] extract_text ✓  {len(texts)} 条 → {out}  字段=\"{field_label}\"",
                        flush=True,
                    )
                else:
                    print(f"[recorder] ⚠️  WARNING: 0 items matched selector {repr(target)}", flush=True)
                    print(f"[recorder]    The selector may be wrong or content not yet rendered.", flush=True)
                    print(f"[recorder]    Try: dom_inspect on a parent container to see real DOM structure.", flush=True)
                    error = (f"⚠️ 提取到 0 条内容。选择器 {repr(target)} 可能不匹配当前页面的真实 DOM 结构。"
                             f"\n建议：先用 dom_inspect 检查父容器的真实子元素结构，再修正选择器。"
                             f"\n / 0 items extracted. Selector {repr(target)} may not match the real DOM structure."
                             f"\nTip: use dom_inspect on a parent container to see actual child elements, then fix the selector.")

                # Generated code：调用 _write_kv_field 写 kv 文件（与录制时完全同构）
                lim_code  = str(limit) if limit else "9999"
                field_lit = repr(
                    (data.get("field") or data.get("field_name") or context or f"步骤 {step_n}").strip()
                    or "extract"
                )
                first_repr = repr(first_for_name)
                body_lines = [
                    f'_sel = {repr(target)}',
                    f'_lim = {lim_code}',
                    'await _wait_spa_ready_for_vision(page, "", timeout_ms=CONFIG["extract_ready_timeout_ms"])',
                    'await _wait_for_content(page, _sel)',
                    '_texts = await page.evaluate(_EXTRACT_JS, [_sel, _lim])',
                    f'_out = CONFIG["tmp_dir"] / {repr(filename)}',
                    f'_field = {field_lit}',
                    f'_write_kv_field(_out, _field, _texts, first_write={first_repr})',
                    'print(f"已提取 {{len(_texts)}} 条 → {{_out}}  字段\\"{_field}\\"")',
                ]
                code_block = _step_code(step_n, context, body_lines)

                # 摘要：结构化信息供 AI 写 python_snippet，不含实际值
                extract_summary = {
                    "file": filename,
                    "field": field_label,
                    "count": len(texts),
                    "read_expr": f'_parse_field(CONFIG["tmp_dir"] / "{filename}", "{field_label}")',
                }

        elif action == "extract_by_vision":
            # ── 视觉识别提取：截图 → 调用视觉 LLM → 写 kv 文件（与 extract_text 同格式）──
            fields      = data.get("fields") or []
            filename    = (value or "output.txt").strip()
            model_key   = (data.get("model_key") or "qwen").strip()
            api_key     = (data.get("api_key") or "").strip()
            crop_sel    = (data.get("crop_selector") or "").strip()

            # 优先用 action 里的 api_key；其次取本次 session 缓存；最后取本地持久化
            if not api_key:
                api_key = _VISION_SESSION.get("api_key", "")
            if not api_key:
                api_key = _load_cached_vision_key(model_key)

            if not fields:
                error = "extract_by_vision 需要 fields 列表 / requires 'fields' list"
            elif model_key not in _VISION_MODELS:
                error = f"未知视觉模型 {model_key!r}，可选：{list(_VISION_MODELS)}"
            elif not api_key:
                error = (
                    f"extract_by_vision 需要 api_key。\n"
                    f"请在 action JSON 里加 \"api_key\": \"sk-xxx\"，\n"
                    f"或先通过 SKILL 对话流程设置 {_VISION_MODELS[model_key]['key_env']}。\n"
                    f"获取地址：{_VISION_MODELS[model_key]['key_url']}"
                )
            else:
                # 0. 截图前等待 SPA 主内容（避免骨架屏 / 未加载完就送视觉 API）
                _vready_ms = int(data.get("vision_ready_timeout_ms") or 45_000)
                await _wait_spa_ready_for_vision(page, crop_sel, timeout_ms=_vready_ms)

                # 1. 截图（可裁剪到指定容器）
                shot_path = shots_dir / f"vision_step_{step_n:02d}_{datetime.now().strftime('%H%M%S')}.png"
                if crop_sel:
                    try:
                        elem = page.locator(crop_sel).first
                        await elem.screenshot(path=str(shot_path))
                    except Exception:
                        await page.screenshot(path=str(shot_path), full_page=False)
                else:
                    await page.screenshot(path=str(shot_path), full_page=False)

                # 2. 调用视觉 API（录制时真实调用，验证字段能提取到）
                try:
                    extracted: dict[str, str] = await _call_vision_api(
                        shot_path.read_bytes(), fields, model_key, api_key
                    )
                except Exception as exc:
                    error = (
                        f"视觉 API 调用失败：{exc}\n"
                        f"请检查：① API Key 是否有效；"
                        f"② 截图是否包含目标内容（可先 scroll 再重试）；"
                        f"③ 网络是否可访问 {_VISION_MODELS[model_key]['base_url']}"
                    )
                    extracted = {}

                if not error:
                    # 3. 写 kv 文件（与 extract_text 完全相同的格式与卡口）
                    first_for_name = filename not in _EXTRACT_OUTPUT_FILES
                    _EXTRACT_OUTPUT_FILES.add(filename)
                    out = _TASK_TMP_DIR / filename

                    for i, field in enumerate(fields):
                        val = str(extracted.get(field, "")).strip()
                        _write_kv_field(out, field, [val] if val else [], first_write=(i == 0 and first_for_name))
                        _EXTRACT_FIELD_REGISTRY.setdefault(filename, []).append(field)

                    # 4. 持久化 key 到本地缓存，下次录制自动复用
                    _save_vision_key(model_key, api_key)
                    _VISION_SESSION["model_key"] = model_key
                    _VISION_SESSION["api_key"]   = api_key

                    # 5. 记录本次视觉步骤（用于 vision_setup.md）
                    _VISION_STEPS.append({
                        "step":    step_n,
                        "fields":  fields,
                        "file":    filename,
                        "model":   _VISION_MODELS[model_key]["model"],
                        "preview": {k: (v[:40] + "…" if len(v) > 40 else v) for k, v in extracted.items()},
                    })

                    print(
                        f"[recorder] extract_by_vision ✓  {len(fields)} 字段 → {out}  "
                        f"model={_VISION_MODELS[model_key]['model']}",
                        flush=True,
                    )
                    for fld, val in extracted.items():
                        print(f"[recorder]   {fld}: {repr(val[:60])}", flush=True)

                    # 摘要（与 extract_text 相同结构）
                    extract_summary = {
                        "file":    filename,
                        "fields":  fields,
                        "model":   _VISION_MODELS[model_key]["model"],
                        "preview": _VISION_STEPS[-1]["preview"],
                        "read_expr": " | ".join(
                            f'_parse_field(CONFIG["tmp_dir"] / "{filename}", "{f}")'
                            for f in fields
                        ),
                    }

                    # 6. 生成代码片段（写入最终脚本）
                    cfg_v = _VISION_MODELS[model_key]
                    body_lines = [
                        f'# 视觉截图前等待页面就绪（骨架屏消失 / 大图或正文出现）',
                        f'await _wait_spa_ready_for_vision(page, {repr(crop_sel)}, '
                        f'timeout_ms=CONFIG["vision_ready_timeout_ms"])',
                        f'# 截图{"（裁剪至 " + crop_sel + "）" if crop_sel else ""}',
                        f'_vision_shot = CONFIG["tmp_dir"] / "vision_step_{step_n:02d}.png"',
                    ]
                    if crop_sel:
                        body_lines += [
                            f'try:',
                            f'    _elem = page.locator({repr(crop_sel)}).first',
                            f'    await _elem.screenshot(path=str(_vision_shot))',
                            f'except Exception:',
                            f'    await page.screenshot(path=str(_vision_shot), full_page=False)',
                        ]
                    else:
                        body_lines.append(
                            'await page.screenshot(path=str(_vision_shot), full_page=False)'
                        )
                    body_lines += [
                        f'_extracted = await _vision_call({repr(fields)}, _vision_shot.read_bytes(), CONFIG)',
                        f'_out = CONFIG["tmp_dir"] / {repr(filename)}',
                        f'_first_v = {repr(first_for_name)}',
                        f'for _i, _f in enumerate({repr(fields)}):',
                        '    _val = str(_extracted.get(_f, "")).strip()',
                        '    _write_kv_field(_out, _f, [_val] if _val else [], first_write=(_i == 0 and _first_v))',
                        'print(f"视觉提取完成 → {_out}  字段：{list(_extracted.keys())}")',
                    ]
                    code_block = _step_code(step_n, context, body_lines)

        elif action == "wait":
            ms = int(value) if value else 2000
            await page.wait_for_timeout(ms)
            code_block = _step_code(step_n, context, [
                f'await page.wait_for_timeout({ms})',
            ])

        elif action == "scroll":
            px = int(value) if value else 500
            try:
                await page.wait_for_load_state("domcontentloaded", timeout=10_000)
            except Exception:
                pass
            vp = page.viewport_size
            if vp:
                await page.mouse.move(vp["width"] // 2, vp["height"] // 2)
            else:
                await page.mouse.move(720, 450)
            await page.mouse.wheel(0, float(px))
            await page.wait_for_timeout(600)   # wait for lazy-load trigger
            code_block = _step_code(step_n, context, [
                f"await _scroll_window(page, {px})",
                "await page.wait_for_timeout(600)",
            ])

        elif action == "scroll_to":
            # Scroll a specific element into view — triggers lazy-load for that section
            await page.evaluate(
                """(sel) => {
                    const el = document.querySelector(sel);
                    if (el) el.scrollIntoView({ behavior: 'smooth', block: 'center' });
                }""",
                target,
            )
            await page.wait_for_timeout(1200)  # wait for lazy content to render
            # Use single-quoted JS string in generated code to avoid escape hell
            js_str = "(s)=>{const e=document.querySelector(s);if(e)e.scrollIntoView({block:'center'})}"
            code_block = _step_code(step_n, context, [
                f'await page.evaluate({repr(js_str)}, {repr(target)})',
                'await page.wait_for_timeout(1200)',
            ])

        elif action == "api_call":
            url = _build_api_url_for_record(data)
            method = (data.get("method") or "GET").upper()
            timeout = float(data.get("timeout") or 60.0)
            save_to = data.get("save_response_to")
            headers = _headers_for_record(data)
            body = data.get("body")
            # verify_ssl：用户显式传 false → 跳过；
            # 未传时自动检测：主机名含下划线（SSL 规范不允许）→ 自动关闭校验
            _host = urllib.parse.urlparse(url).hostname or ""
            _has_underscore_host = "_" in _host
            verify_ssl = data.get("verify_ssl", not _has_underscore_host)
            if _has_underscore_host and data.get("verify_ssl") is None:
                print(f"[recorder] ⚠️  主机名含下划线（{_host}），已自动关闭 SSL 校验（verify=False）"
                      f" / hostname contains underscore ({_host}), SSL verification auto-disabled (verify=False)", flush=True)
            try:
                async with httpx.AsyncClient(timeout=timeout, verify=verify_ssl) as _hc:
                    if method == "GET":
                        r = await _hc.get(url, headers=headers)
                    elif method == "POST":
                        if isinstance(body, dict):
                            r = await _hc.post(url, json=body, headers=headers)
                        else:
                            r = await _hc.post(
                                url,
                                content=body if body is not None else "",
                                headers=headers,
                            )
                    else:
                        r = await _hc.request(method, url, headers=headers)
                    r.raise_for_status()
                    _api_text = r.text
            except httpx.ConnectError as _ssl_exc:
                _msg = str(_ssl_exc)
                if "SSL" in _msg or "certificate" in _msg.lower() or "CERTIFICATE" in _msg:
                    raise RuntimeError(
                        f"SSL 证书验证失败（{_msg}）。\n"
                        "若目标服务使用了含下划线的主机名或自签名证书，可在 api_call 步骤中加上 "
                        "\"verify_ssl\": false 跳过校验（仅用于测试环境）。\n"
                        f" / SSL certificate verification failed ({_msg}).\n"
                        "If the host has an underscore in its name or uses a self-signed cert, "
                        "add \"verify_ssl\": false to the api_call step (test environments only)."
                    ) from _ssl_exc
                raise
            if save_to:
                _out = _TASK_TMP_DIR / save_to
                _out.parent.mkdir(parents=True, exist_ok=True)
                _out.write_text(_api_text, encoding="utf-8")
            code_block = _step_code(step_n, context, _api_codegen_body(context, data))

        elif action == "excel_write":
            err_ex = _excel_write_run(data)
            if err_ex:
                error = err_ex
            else:
                code_block = _step_code(step_n, context, _excel_write_codegen_lines(data))

        elif action == "word_write":
            err_wd = _word_write_run(data)
            if err_wd:
                error = err_wd
            else:
                code_block = _step_code(step_n, context, _word_write_codegen_lines(data))

        elif action == "python_snippet":
            # Execute the snippet NOW (same as api_call / excel_write run at record time).
            # This validates dependencies, file existence, and logic before writing code_block.
            raw_code = (data.get("code") or "").rstrip()
            if not raw_code.strip():
                error = "python_snippet 需要非空的 code 字段 / python_snippet requires a non-empty 'code' field"
            else:
                error = await _python_snippet_run(raw_code, page)
                if not error:
                    # 若本步成功使用了 page.evaluate，标记 session 级别豁免标志
                    # Mark session-level exemption when page.evaluate is used successfully
                    if "page.evaluate" in raw_code:
                        global _SESSION_HAS_PAGE_EVALUATE
                        _SESSION_HAS_PAGE_EVALUATE = True
                    # 结构性卡口：验证代码通过 _parse_field 读取提取文件（_SESSION_HAS_PAGE_EVALUATE 时内部豁免）
                    # Structural gate (internally exempted when _SESSION_HAS_PAGE_EVALUATE is True)
                    error = _check_snippet_reads_extract_files(raw_code)
                if not error:
                    code_block = _step_code(step_n, context, raw_code.splitlines())

        elif action == "merge_files":
            # Pure file operation: read sources from Desktop, concatenate, write target.
            # Does NOT interact with the browser.
            sources   = data.get("sources") or []  # list of Desktop filenames
            target_fn = data.get("target") or data.get("value") or ""
            separator = data.get("separator", "\n\n")
            if not sources or not target_fn:
                error = "merge_files 需要 sources（列表）和 target（目标文件名）/ merge_files requires 'sources' (list) and 'target' (output filename)"
            else:
                parts: list[str] = []
                for src in sources:
                    # 源文件优先在 tmp_dir 查找（extract/api 中间文件），其次 output_dir
                    p = _TASK_TMP_DIR / src
                    if not p.exists():
                        p = Path.home() / "Desktop" / src
                    if p.exists():
                        parts.append(p.read_text(encoding="utf-8"))
                    else:
                        print(f"[recorder] ⚠️  merge_files：文件不存在，跳过 / file not found, skipping: {src}", flush=True)
                out_path = Path.home() / "Desktop" / target_fn
                out_path.write_text(separator.join(parts), encoding="utf-8")
                print(f"[recorder] merge_files → {out_path}（{len(parts)}/{len(sources)} 个源文件 / source files merged）", flush=True)
            # Code generation（源文件从 tmp_dir 读，合并结果写到 output_dir）
            sep_repr   = repr(separator)
            srcs_repr  = repr(sources)
            tgt_repr   = repr(target_fn)
            body_lines = [
                f"_merge_sources = {srcs_repr}",
                f"_merge_sep = {sep_repr}",
                "_merge_parts = []",
                "for _src in _merge_sources:",
                '    _p = CONFIG["tmp_dir"] / _src',
                '    if not _p.exists():',
                '        _p = CONFIG["output_dir"] / _src',
                "    if _p.exists():",
                '        _merge_parts.append(_p.read_text(encoding="utf-8"))',
                '    else:',
                '        print(f"⚠️  merge_files：文件不存在，跳过 {_src}")',
                f'(CONFIG["output_dir"] / {tgt_repr}).write_text(_merge_sep.join(_merge_parts), encoding="utf-8")',
                f'print("已合并到", CONFIG["output_dir"] / {tgt_repr})',
            ]
            code_block = _step_code(step_n, context, body_lines)

        elif action == "snapshot":
            pass  # read-only DOM inspection — NOT logged to script

        elif action == "dom_inspect":
            # Diagnostic: return child structure of a container element.
            # NOT logged to the script — only used to discover real selectors.
            _di_ms = int(data.get("extract_ready_timeout_ms") or 25_000)
            await _wait_spa_ready_for_vision(page, str(target or "").strip(), timeout_ms=_di_ms)
            result = await page.evaluate("""(sel) => {
                const el = document.querySelector(sel);
                if (!el) return { found: false, message: 'Element not found: ' + sel };
                const children = Array.from(el.querySelectorAll('*'))
                    .slice(0, 50)
                    .map(c => ({
                        tag:    c.tagName.toLowerCase(),
                        id:     c.id || null,
                        testid: c.getAttribute('data-testid') || null,
                        cls:    Array.from(c.classList).slice(0, 2).join(' ') || null,
                        aria:   c.getAttribute('aria-label') || null,
                        text:   (c.textContent || '').replace(/\\s+/g, ' ').trim().slice(0, 60),
                    }));
                return { found: true, outerTag: el.tagName.toLowerCase(), children };
            }""", target)
            # Attach inspect result directly to action result (not to code_block)
            if isinstance(result, dict) and result.get("found"):
                children = result.get("children", [])
                # Print structured summary for the LLM to read
                lines = [f"[dom_inspect] 容器 {repr(target)} 共 {len(children)} 个子元素 / container {repr(target)} has {len(children)} children:"]
                for c in children[:30]:
                    sel_hint = (
                        f"#{c['id']}" if c['id']
                        else f"[data-testid=\"{c['testid']}\"]" if c['testid']
                        else f"[aria-label=\"{c['aria']}\"]" if c['aria']
                        else f".{c['cls'].split()[0]}" if c['cls']
                        else c['tag']
                    )
                    print(f"  {sel_hint}  「{c['text'][:50]}」", flush=True)
                # Return children in result for rpa_manager to display
                code_block = None  # not logged
                inspect_children = children
            else:
                error = result.get("message", "dom_inspect failed") if isinstance(result, dict) else "dom_inspect failed"

        else:
            error = f"未知 action / unknown action: {action!r}"

    except Exception as exc:
        error = str(exc)

    # Screenshot after every action (proof + visual feedback for user)
    shot_path: "Path | None" = None
    snap: list = []
    sections: list = []
    data_groups: list = []
    url = ""

    if page is not None:
        # snapshot：先等主内容再截图/采 DOM，减少把骨架屏当成「真实结构」
        if action == "snapshot":
            try:
                _sn_ms = int(data.get("extract_ready_timeout_ms") or 25_000)
                await _wait_spa_ready_for_vision(page, "", timeout_ms=_sn_ms)
            except Exception:
                pass
        ts        = datetime.now().strftime("%H%M%S")
        label     = "snapshot" if action == "snapshot" else f"step_{step_n:02d}"
        shot_path = shots_dir / f"{label}_{ts}.png"
        try:
            await page.screenshot(path=str(shot_path), full_page=False)
        except Exception:
            shot_path = None

        # Run DOM snapshot — always includes all three layers (items / sections / data_groups)
        raw_snap = await _snapshot(page)
        snap        = raw_snap.get("items", [])
        sections    = raw_snap.get("sections", [])
        data_groups = raw_snap.get("data_groups", [])

        url = page.url

    out = {
        "success":      error is None,
        "error":        error,
        "code_block":   code_block,
        "screenshot":   str(shot_path) if shot_path else None,
        "url":          url,
        "snapshot":     snap,
        "sections":     sections,
        "data_groups":  data_groups,
    }
    if inspect_children is not None:
        out["_inspect_children"] = inspect_children
    if extract_summary is not None:
        out["extract_summary"] = extract_summary
    return out


# ── Script builder ───────────────────────────────────────────────────────────

def _build_final_script(
    task_name: str,
    code_blocks: list[str],
    *,
    use_openpyxl: bool = False,
    use_docx: bool = False,
    cookies_file: str = "",
    vision_session: Optional[dict] = None,
    cdp_url: str = "",
) -> str:
    ts    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    _task_slug = _slugify_for_path(task_name)  # 用于生成脚本的 tmp_dir 路径字面量
    steps = "\n\n".join(code_blocks) if code_blocks else "            pass  # 无录制步骤"
    kv_write_src   = inspect.getsource(_write_kv_field)
    parse_field_src = inspect.getsource(_parse_field)
    extract_js_repr = repr(_EXTRACT_JS_MIN)

    # SPA 主内容就绪等待（extract_text / snapshot / dom_inspect / 视觉截图 共用）
    spa_ready_src = inspect.getsource(_wait_spa_ready_for_vision) + "\n\n"

    # ── 视觉识别 helper（仅当本次录制有 extract_by_vision 步骤时插入）──────────
    vision_config_block = ""
    vision_helper_block = ""
    if vision_session and vision_session.get("api_key"):
        vs = vision_session
        mk = vs.get("model_key", "qwen")
        vcfg = _VISION_MODELS.get(mk, _VISION_MODELS["qwen"])
        # 块末必须换行：否则外层 f-string 的「}}」会变成紧贴 # 注释的「}」，整颗括号被注释掉，CONFIG 语法错误
        vision_config_block = f"""\
    # ── 视觉识别配置（由 OpenClaw RPA 录制时自动生成）──────────────────────────
    "vision_model":    {repr(vcfg["model"])},
    "vision_base_url": {repr(vcfg["base_url"])},
    "vision_api_key":  {repr(vs["api_key"])},   # {vcfg["key_env"]}
    "vision_ready_timeout_ms": 45_000,  # 视觉截图前等待；Airbnb 等可调 60_000
"""
        vision_helper_block = """\

async def _vision_call(fields: list, image_bytes: bytes, cfg: dict) -> dict:
    \"\"\"调用视觉 LLM 从截图提取字段。模型/Key 由 CONFIG 注入，无需手动配置。\"\"\"
    import base64 as _vb64, re as _vre
    b64 = _vb64.b64encode(image_bytes).decode()
    tmpl = {f: "" for f in fields}
    prompt = (
        f"从截图中提取以下字段，只返回 JSON，不要解释：\\n"
        + __import__("json").dumps(tmpl, ensure_ascii=False)
        + "\\n规则：①只提取可见文字；②看不到设为空字符串；③价格保留原始格式。"
    )
    payload = {
        "model": cfg["vision_model"],
        "messages": [{"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            {"type": "text", "text": prompt},
        ]}],
        "max_tokens": 500,
    }
    base = cfg["vision_base_url"].rstrip("/")
    _vt = httpx.Timeout(300.0, connect=30.0)
    async with httpx.AsyncClient(timeout=_vt, verify=False) as _hc:
        _r = await _hc.post(
            f"{base}/chat/completions",
            headers={"Authorization": f"Bearer {cfg['vision_api_key']}", "Content-Type": "application/json"},
            json=payload,
        )
        _r.raise_for_status()
    raw = _r.json()["choices"][0]["message"]["content"].strip()
    raw = _vre.sub(r"^```(?:json)?\\s*|\\s*```$", "", raw, flags=_vre.MULTILINE).strip()
    return __import__("json").loads(raw)
"""

    # Collect env vars used by __ENV:VAR__ placeholders in all api_call steps.
    # They appear in code blocks as:  os.environ.get("VAR_NAME", "")
    _env_var_re = re.compile(r'os\.environ\.get\("([A-Za-z_][A-Za-z0-9_]*)",\s*""\)')
    env_vars = sorted(set(_env_var_re.findall(steps)))

    # Replace inline os.environ.get(...) with CONFIG["VAR"] so keys are centralised.
    for _v in env_vars:
        steps = steps.replace(f'os.environ.get("{_v}", "")', f'CONFIG["{_v}"]')

    # Build CONFIG entries for env vars (one line each, with export hint).
    if env_vars:
        env_config_lines = "\n    # --- API 密钥（从环境变量读取；个人使用时也可直接在此填写）"
        for _v in env_vars:
            env_config_lines += f'\n    "{_v}": os.environ.get("{_v}", ""),'
        env_config_lines += "\n"
    else:
        env_config_lines = ""

    # Startup check: fail fast with helpful message if any required key is empty.
    if env_vars:
        env_list_repr = repr(env_vars)
        startup_check = f"""\
_missing = [v for v in {env_list_repr} if not CONFIG.get(v)]
if _missing:
    print("\\n❌  以下 API 密钥未配置，脚本无法运行 / The following API keys are not set, cannot run:")
    for _v in _missing:
        print(f"   {{_v}} 为空 / is empty  →  export {{_v}}='your-key'")
        print(f"   （或在脚本 CONFIG 字典里直接填写 / or set it directly in the CONFIG dict above）")
    raise SystemExit(1)

"""
    else:
        startup_check = ""

    office_imports = ""
    if use_openpyxl:
        office_imports += (
            "\nfrom openpyxl import Workbook, load_workbook\n"
            "from openpyxl.utils import get_column_letter\n"
        )
    if use_docx:
        office_imports += "\nfrom docx import Document\n"

    cookies_path_line = (
        f'\n    # 已保存的登录 Cookie 路径（由 #rpa-login 生成；留空则不注入）\n'
        f'    "cookies_path":  {repr(cookies_file)},'
        if cookies_file else
        '\n    # 若有已保存的登录 Cookie，填写路径（由 rpa_manager login-start 生成）\n'
        '    "cookies_path":  "",'
    )

    _cdp_run_block = ""
    if cdp_url:
        # CDP 模式：通过 CDP 连接已有 Chrome
        CDP_URL_REPR = repr(cdp_url)
        _cdp_run_block = f"""\
async def run():
    async with async_playwright() as p:
        _cdp_url = os.environ.get("CDP_URL") or {CDP_URL_REPR}
        print(f"通过 CDP 连接已有 Chrome：{{_cdp_url}}")
        browser = await p.chromium.connect_over_cdp(_cdp_url)
        contexts = browser.contexts
        if contexts:
            context = contexts[0]
        else:
            context = await browser.new_context()
        # 若配置了 cookies_path，注入 Cookie 模拟已登录态（由 rpa_manager login-start 生成）
        import json as _json
        _cp = CONFIG.get("cookies_path", "")
        if _cp and Path(_cp).exists():
            await context.add_cookies(_json.loads(Path(_cp).read_text()))
        page = await context.new_page()
        page.set_default_timeout(CONFIG["timeout"])

        try:
{{steps}}

        except PlaywrightTimeout as e:
            await page.screenshot(path="error_timeout.png")
            raise RuntimeError(f"超时：{{e}}") from e
        except Exception:
            await page.screenshot(path="error_unexpected.png")
            raise
        finally:"""
    else:
        # 本地模式：启动新浏览器
        _cdp_run_block = """\
async def run():
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=CONFIG["headless"],
            args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
            slow_mo=CONFIG["slow_mo"],
        )
        context = await browser.new_context(
            user_agent=_UA,
            viewport={{"width": 1440, "height": 900}},
            locale="en-US",
            timezone_id="America/New_York",
            extra_http_headers={{"Accept-Language": "en-US,en;q=0.9"}},
        )
        await context.add_init_script(
            "Object.defineProperty(navigator, 'webdriver', {{get: () => undefined}})"
        )
        # 若配置了 cookies_path，注入 Cookie 模拟已登录态（由 rpa_manager login-start 生成）
        import json as _json
        _cp = CONFIG.get("cookies_path", "")
        if _cp and Path(_cp).exists():
            await context.add_cookies(_json.loads(Path(_cp).read_text()))
        page = await context.new_page()
        page.set_default_timeout(CONFIG["timeout"])

        try:
{{steps}}

        except PlaywrightTimeout as e:
            await page.screenshot(path="error_timeout.png")
            raise RuntimeError(f"超时：{{e}}") from e
        except Exception:
            await page.screenshot(path="error_unexpected.png")
            raise
        finally:"""

    return f"""\
# pip install playwright httpx openpyxl python-docx && playwright install chromium
# 任务：{task_name}
# 录制时间：{ts}
# 由 OpenClaw RPA Recorder（headed 真实录制）生成 — 可脱离 OpenClaw 独立运行

import asyncio
import datetime
import json
import os
import re
import urllib.parse
from pathlib import Path

import httpx{office_imports}
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeout

CONFIG = {{
    "output_dir":    Path.home() / "Desktop",
    # 中间临时文件目录（提取结果 / API 响应 / 聚合 JSON）——独立于最终产物，避免跨任务污染
    "tmp_dir":       Path("/tmp") / {repr(_task_slug)},
    "headless":      False,
    "timeout":       60_000,
    "slow_mo":       300,
    # 导航后等待 SPA 内容渲染的额外时间（重型 SPA 如 Yahoo Finance 需要 1-2 秒）
    "spa_settle_ms": 1_500,
    # extract_text 等待目标元素出现的超时（毫秒）
    "content_wait":  15_000,
    # httpx 调用外部 API 的超时（秒）
    "api_timeout":   60.0,
    # extract_text / snapshot / dom_inspect 前等待主内容（骨架屏后再读 DOM）
    "extract_ready_timeout_ms": 30_000,{cookies_path_line}{env_config_lines}{vision_config_block}
}}
# 确保临时目录存在
CONFIG["tmp_dir"].mkdir(parents=True, exist_ok=True)

{startup_check}_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

{kv_write_src}

{parse_field_src}

_EXTRACT_JS = {extract_js_repr}
{spa_ready_src}{vision_helper_block}


async def _wait_for_content(page, selector: str) -> None:
    \"\"\"等待 selector 对应的元素出现在 DOM 中（容错：超时也继续）。\"\"\"
    try:
        await page.wait_for_selector(selector, timeout=CONFIG["content_wait"])
    except Exception:
        pass  # 元素未出现也继续，evaluate 会返回空列表


async def _scroll_window(page, dy: int) -> None:
    \"\"\"窗口滚动：导航后若再用 evaluate(scrollBy)，易因执行上下文销毁报错；用 mouse.wheel 并在滚动前等待页面稳定。\"\"\"
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=10_000)
    except Exception:
        pass
    vp = page.viewport_size
    if vp:
        await page.mouse.move(vp["width"] // 2, vp["height"] // 2)
    else:
        await page.mouse.move(720, 450)
    await page.mouse.wheel(0, float(dy))


{_cdp_run_block}
            await browser.close()


if __name__ == "__main__":
    asyncio.run(run())
"""


# ── Login capture mode / 登录捕获模式 ────────────────────────────────────────

async def login_capture_main():
    """登录捕获模式：打开浏览器到登录页，等待用户完成登录，再导出 Cookie 保存到文件。
    Login capture mode: open browser to login page, wait for user to finish login,
    then export cookies to a file for later injection."""
    SESSION_DIR.mkdir(parents=True, exist_ok=True)
    pid_path = SESSION_DIR / "server.pid"
    pid_path.write_text(str(os.getpid()))

    task_data         = json.loads((SESSION_DIR / "task.json").read_text())
    login_url         = task_data["login_url"]
    cookies_output    = task_data.get("cookies_output", str(SESSION_DIR / "cookies.json"))
    cookies_meta_out  = task_data.get("cookies_meta_output", str(SESSION_DIR / "cookies_meta.json"))

    from playwright.async_api import async_playwright

    _cdp_url = task_data.get("cdp_url", "")

    async with async_playwright() as p:
        if _cdp_url:
            print(f"[login_capture] 通过 CDP 连接已有 Chrome：{_cdp_url}", flush=True)
            browser = await p.chromium.connect_over_cdp(_cdp_url)
            contexts = browser.contexts
            if contexts:
                ctx = contexts[0]
            else:
                ctx = await browser.new_context()
        else:
            browser = await p.chromium.launch(
                headless=False,
                args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
                slow_mo=0,
            )
            ctx = await browser.new_context(
                user_agent=_UA,
                viewport={"width": 1440, "height": 900},
                locale="zh-CN",
                timezone_id="Asia/Shanghai",
                extra_http_headers={"Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8"},
            )
            await ctx.add_init_script(
                "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"
            )
        page = await ctx.new_page()
        page.set_default_timeout(120_000)

        await page.goto(login_url)

        # 通知 rpa_manager 浏览器已就绪 / Signal rpa_manager that the browser is ready
        (SESSION_DIR / "ready").write_text("1")
        print(f"[login_capture] ready — {login_url}", flush=True)

        # 轮询等待 login_done 指令 / Poll for login_done command from rpa_manager
        cmd_path = SESSION_DIR / "cmd.json"
        last_seq = -1
        while True:
            if cmd_path.exists():
                try:
                    data = json.loads(cmd_path.read_text())
                    seq  = data.get("seq", 0)
                    if seq > last_seq:
                        last_seq = seq
                        if data.get("action") == "login_done":
                            break
                except Exception:
                    pass
            await asyncio.sleep(POLL_INTERVAL)

        # 导出当前 context 的全部 Cookie / Export all cookies from the current browser context
        cookies = await ctx.cookies()
        Path(cookies_output).parent.mkdir(parents=True, exist_ok=True)
        Path(cookies_output).write_text(
            json.dumps(cookies, indent=2, ensure_ascii=False), encoding="utf-8"
        )

        # 生成 meta 摘要 / Build metadata summary (total, session-type count, earliest expiry)
        expiries = [c["expires"] for c in cookies if c.get("expires", -1) > 0]
        earliest = min(expiries) if expiries else None
        meta = {
            "saved_at":       datetime.now().isoformat(timespec="seconds"),
            "total":          len(cookies),
            "session_cookies": sum(1 for c in cookies if c.get("expires", -1) <= 0),
            "earliest_expires": (
                datetime.fromtimestamp(earliest).isoformat(timespec="seconds")
                if earliest else None
            ),
            "hint": "实际有效期以服务端策略为准，可能早于参考时间",
        }
        Path(cookies_meta_out).write_text(
            json.dumps(meta, indent=2, ensure_ascii=False), encoding="utf-8"
        )

        await browser.close()

        # 写 login_done 标记通知 rpa_manager / Write done marker so rpa_manager unblocks
        (SESSION_DIR / "login_done").write_text("1")
    pid_path.unlink(missing_ok=True)
    print(f"[login_capture] done — {len(cookies)} cookies → {cookies_output}", flush=True)


# ── Server main loop ─────────────────────────────────────────────────────────

async def server_main():
    SESSION_DIR.mkdir(parents=True, exist_ok=True)
    pid_path = SESSION_DIR / "server.pid"
    pid_path.write_text(str(os.getpid()))

    _reset_extract_output_tracking()

    task_data = json.loads((SESSION_DIR / "task.json").read_text())
    task_name = task_data["task"]
    shots_dir = SESSION_DIR / "screenshots"
    shots_dir.mkdir(exist_ok=True)

    # 初始化本 session 的临时文件目录（/tmp/{task_slug}/，隔离不同任务的中间文件）
    global _TASK_TMP_DIR
    _task_slug = _slugify_for_path(task_name)
    _TASK_TMP_DIR = Path("/tmp") / _task_slug
    _TASK_TMP_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[recorder] tmp_dir → {_TASK_TMP_DIR}", flush=True)

    code_blocks: list[str] = []
    step_n = 0
    use_openpyxl = False
    use_docx = False

    needs_browser = task_data.get("needs_browser", True)

    if not needs_browser:
        # File/API-only mode (B/C/F/N): no Playwright browser needed.
        # 纯文件/API 模式（B/C/F/N）：不启动浏览器。
        (SESSION_DIR / "ready").write_text("1")
        print(f"[recorder] ready (no-browser mode) — task: {task_name}", flush=True)

        last_seq = -1
        cmd_path = SESSION_DIR / "cmd.json"

        while True:
            if cmd_path.exists():
                try:
                    data = json.loads(cmd_path.read_text())
                    seq  = data.get("seq", 0)
                    if seq > last_seq:
                        last_seq = seq
                        action   = data.get("action", "")

                        if action == "shutdown":
                            break

                        if action == "excel_write":
                            use_openpyxl = True
                        if action == "word_write":
                            use_docx = True
                        if action == "python_snippet":
                            snippet_code = (data.get("code") or "")
                            if "load_workbook" in snippet_code or "openpyxl" in snippet_code or "Workbook" in snippet_code:
                                use_openpyxl = True
                            if "Document" in snippet_code or "from docx" in snippet_code:
                                use_docx = True

                        if action != "snapshot":
                            step_n += 1

                        result = await _do_action(None, data, step_n, shots_dir)

                        if result.get("code_block"):
                            code_blocks.append(result["code_block"])

                        (SESSION_DIR / f"result_{seq}.json").write_text(
                            json.dumps(result, ensure_ascii=False, indent=2)
                        )
                except Exception as exc:
                    try:
                        (SESSION_DIR / f"result_{last_seq}.json").write_text(
                            json.dumps(
                                {"success": False, "error": str(exc),
                                 "code_block": None, "snapshot": []},
                                ensure_ascii=False,
                            )
                        )
                    except Exception:
                        pass

            await asyncio.sleep(POLL_INTERVAL)

    else:
        # Browser mode (A/D/E/G): launch headed Chromium or connect via CDP.
        # 浏览器模式（A/D/E/G）：启动有界面 Chromium，或通过 CDP 连接已有 Chrome。
        from playwright.async_api import async_playwright

        _cdp_url = task_data.get("cdp_url", "")

        async with async_playwright() as p:
            if _cdp_url:
                print(f"[recorder] 通过 CDP 连接已有 Chrome：{_cdp_url}", flush=True)
                browser = await p.chromium.connect_over_cdp(_cdp_url)
                # 使用默认浏览器上下文（保留用户已有的登录态）
                contexts = browser.contexts
                if contexts:
                    ctx = contexts[0]
                else:
                    ctx = await browser.new_context()
                use_active_tab = task_data.get("use_active_tab", False)
                if use_active_tab and ctx.pages:
                    # 复用当前已激活的标签页
                    page = ctx.pages[-1]
                    print(f"[recorder] 使用已有标签页：{page.url}", flush=True)
                else:
                    page = await ctx.new_page()
            else:
                browser = await p.chromium.launch(
                    headless=False,
                    args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
                    slow_mo=200,
                )
                ctx = await browser.new_context(
                    user_agent=_UA,
                    viewport={"width": 1440, "height": 900},
                    locale="en-US",
                    timezone_id="America/New_York",
                    extra_http_headers={"Accept-Language": "en-US,en;q=0.9"},
                )
                await ctx.add_init_script(
                    "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"
                )
                page = await ctx.new_page()
            page.set_default_timeout(60_000)

            # Signal ready to rpa_manager (it polls for this file)
            (SESSION_DIR / "ready").write_text("1")
            print(f"[recorder] ready — task: {task_name}", flush=True)

            last_seq = -1
            cmd_path = SESSION_DIR / "cmd.json"

            while True:
                if cmd_path.exists():
                    try:
                        data = json.loads(cmd_path.read_text())
                        seq  = data.get("seq", 0)
                        if seq > last_seq:
                            last_seq = seq
                            action   = data.get("action", "")

                            if action == "shutdown":
                                break

                            if action == "excel_write":
                                use_openpyxl = True
                            if action == "word_write":
                                use_docx = True
                            if action == "python_snippet":
                                snippet_code = (data.get("code") or "")
                                if "load_workbook" in snippet_code or "openpyxl" in snippet_code or "Workbook" in snippet_code:
                                    use_openpyxl = True
                                if "Document" in snippet_code or "from docx" in snippet_code:
                                    use_docx = True

                            if action != "snapshot":
                                step_n += 1

                            result = await _do_action(page, data, step_n, shots_dir)

                            if result.get("code_block"):
                                code_blocks.append(result["code_block"])

                            (SESSION_DIR / f"result_{seq}.json").write_text(
                                json.dumps(result, ensure_ascii=False, indent=2)
                            )
                    except Exception as exc:
                        try:
                            (SESSION_DIR / f"result_{last_seq}.json").write_text(
                                json.dumps(
                                    {"success": False, "error": str(exc),
                                     "code_block": None, "snapshot": []},
                                    ensure_ascii=False,
                                )
                            )
                        except Exception:
                            pass

                await asyncio.sleep(POLL_INTERVAL)

            await browser.close()

    # Compile and save final script
    script = _build_final_script(
        task_name,
        code_blocks,
        use_openpyxl=use_openpyxl,
        use_docx=use_docx,
        cookies_file=task_data.get("cookies_file", ""),
        vision_session=dict(_VISION_SESSION) if _VISION_SESSION else None,
        cdp_url=task_data.get("cdp_url", ""),
    )
    (SESSION_DIR / "script_log.py").write_text(script, encoding="utf-8")
    # 若本次录制有视觉步骤，保存步骤信息供 rpa_manager 生成 vision_setup.md
    if _VISION_STEPS:
        (SESSION_DIR / "vision_steps.json").write_text(
            json.dumps({
                "task": task_name,
                "model_key": _VISION_SESSION.get("model_key", "qwen"),
                "model": _VISION_MODELS.get(_VISION_SESSION.get("model_key","qwen"), {}).get("model",""),
                "steps": _VISION_STEPS,
            }, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    (SESSION_DIR / "done").write_text("1")
    pid_path.unlink(missing_ok=True)
    print(f"[recorder] done — {len(code_blocks)} steps — script saved.", flush=True)


if __name__ == "__main__":
    _task_file = SESSION_DIR / "task.json"
    _mode = "record"
    if _task_file.exists():
        try:
            _mode = json.loads(_task_file.read_text()).get("mode", "record")
        except Exception:
            pass
    if _mode == "login_capture":
        asyncio.run(login_capture_main())
    else:
        asyncio.run(server_main())
